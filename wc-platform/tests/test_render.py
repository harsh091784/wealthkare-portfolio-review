"""
tests/test_render.py

Render smoke tests: every fixture must build a real PDF without raising,
and its page count is recorded in the test output so an unexpected jump
is visible on every run.

Page count is deliberately REPORTED, not asserted against a hardcoded
number - pagination legitimately shifts when content changes, and a test
that fails on every copy edit gets muted. What is asserted is that the
build completes, produces a non-trivial PDF, and that the rendered text
contains no banned target-allocation language (invariant 12, checked
here against the actual rendered document rather than the source
strings).
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

import pytest

from pipeline.docx_builder import (
    SECTION_TITLES,
    VERBATIM_ASSET_SECTIONS,
    build_report,
)
from pipeline.mindmap import generate_mindmap
from tests.conftest import FIXTURE_NAMES
from tests.fixtures import ALL_FIXTURE_BUILDERS

FORBIDDEN_RENDERED_COPY = re.compile(r"target\s+(band|allocation)|drift", re.IGNORECASE)

PAGE_COUNTS: dict = {}


def _generated_copy_only(rendered_text: str) -> str:
    """Returns the rendered text with the body of every verbatim asset
    section removed, so what's left is only copy this pipeline generated.

    Scoping note: invariant 12 bans target-allocation/drift language in
    GENERATED report copy. The Director's letter legitimately contains the
    word "drifts" - "the rupee, while it drifts lower over time" - which
    is currency movement, nothing to do with asset allocation. The ban is
    on copy this pipeline writes, not on the Director's prose.

    The exclusion is BY NAMED SECTION (docx_builder.VERBATIM_ASSET_SECTIONS
    = DIRECTORS_MESSAGE, THANK_YOU), not by page offset and not by fuzzy
    matching against the asset files' text. That distinction is the point:
    a fuzzy matcher drops any rendered sentence that happens to resemble
    something in those .docx files, so verbatim content added later - or
    generated copy that coincidentally echoes a phrase from the letters -
    gets silently skipped and the invariant quietly stops covering it.
    Excluding whole named sections means everything outside those two
    headings is scanned, and a new verbatim section is scanned (and fails
    loudly) until it is deliberately added to VERBATIM_ASSET_SECTIONS.

    Section boundaries come from the headings themselves: every section
    heading renders as its own line in the extracted text, while the TOC
    entries carry dot leaders ("Director's Message......4") and so never
    match. A heading line for an excluded section opens a skip; the next
    heading line closes it.
    """
    lines = rendered_text.split("\n")
    heading_lines = {i for i, line in enumerate(lines) if line.strip() in SECTION_TITLES}
    excluded_lines = {
        i for i, line in enumerate(lines) if line.strip() in VERBATIM_ASSET_SECTIONS
    }

    found = {lines[i].strip() for i in excluded_lines}
    missing = set(VERBATIM_ASSET_SECTIONS) - found
    assert not missing, (
        f"verbatim section heading(s) {sorted(missing)} never appeared as a heading line in the "
        f"rendered text. VERBATIM_ASSET_SECTIONS has drifted from the actual rendered headings - "
        f"without a match the section is scanned as generated copy and invariant 12 fails on the "
        f"Director's currency-drift sentence for reasons that have nothing to do with allocation."
    )

    kept, skipping = [], False
    for i, line in enumerate(lines):
        if i in excluded_lines:
            skipping = True
            continue
        if skipping and i in heading_lines:
            skipping = False
        if not skipping:
            kept.append(line)
    return "\n".join(kept)


THINGS_TO_DO_HEADERS = ["#", "Action", "Scheme", "What to do", "Deadline", "Priority"]


def _things_to_do_rows(docx_path: Path) -> list:
    """The Things To Do table's DATA rows, read back out of the rendered
    .docx. Located by its header row rather than by table index, so the
    check doesn't quietly start reading a different table when section
    order changes."""
    from docx import Document
    for table in Document(str(docx_path)).tables:
        header = [c.text.strip() for c in table.rows[0].cells]
        if header == THINGS_TO_DO_HEADERS:
            return table.rows[1:]
    raise AssertionError(
        f"no table with headers {THINGS_TO_DO_HEADERS} found in {docx_path.name} - the "
        f"Things To Do table is missing from the rendered report, or its columns changed."
    )


def _soffice_available() -> bool:
    from pipeline.pdf_converter import _find_soffice
    try:
        _find_soffice()
        return True
    except FileNotFoundError:
        return False


requires_soffice = pytest.mark.skipif(
    not _soffice_available(), reason="LibreOffice (soffice) not available"
)


def _docx_full_text(docx_path: Path) -> str:
    """All user-visible text in a .docx: body paragraphs, table cells, and
    header/footer content. Used to prove the editable deliverable carries
    the same content as the PDF rather than silently degrading."""
    from docx import Document
    doc = Document(str(docx_path))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer):
            chunks += [p.text for p in footer.paragraphs]
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        chunks.append(cell.text)
    return "\n".join(c for c in chunks if c)


@pytest.mark.render
@requires_soffice
@pytest.mark.parametrize("fixture_name", FIXTURE_NAMES)
def test_fixture_produces_both_deliverables(fixture_name, tmp_path, record_property):
    """Every fixture must produce BOTH deliverables without raising - the
    PDF (primary download) and the .docx (secondary, editable) - and the
    PDF's page count is recorded so an unexpected jump is visible.

    Also asserts the two don't diverge: the same generation timestamp and
    the same section headings must be present in both."""
    from pipeline.pdf_converter import convert_docx_to_pdf

    # Each fixture gets its own mind map, built from its own transactions.
    mindmap_path = tmp_path / "mindmap.png"
    builder = ALL_FIXTURE_BUILDERS[fixture_name]
    fixture = builder(mindmap_path=mindmap_path)

    # One timestamp per run, shared by both outputs (as the real
    # orchestrator does) - otherwise docx and PDF could differ by seconds.
    from datetime import datetime
    fixture.ctx.generated_at = datetime(2026, 8, 16, 15, 53)

    if fixture.mindmap_recs:
        generate_mindmap(fixture.mindmap_recs, client_name=fixture.ctx.client_name,
                         output_path=mindmap_path)

    docx_dir = tmp_path / "docx"
    pdf_dir = tmp_path / "pdf"   # must differ from docx dir (bug H9)
    docx_dir.mkdir(parents=True, exist_ok=True)

    # --- Deliverable 2: the editable .docx ---
    docx_path = docx_dir / f"{fixture_name}.docx"
    build_report(fixture.ctx, docx_path)
    assert docx_path.exists(), f"[{fixture_name}] docx deliverable was not produced"
    assert docx_path.stat().st_size > 10_000, (
        f"[{fixture_name}] docx is only {docx_path.stat().st_size} bytes - likely truncated"
    )

    # --- Deliverable 1: the primary PDF, rendered FROM that same docx ---
    pdf_path = convert_docx_to_pdf(docx_path, pdf_dir)
    assert pdf_path.exists(), f"[{fixture_name}] PDF deliverable was not produced"
    assert pdf_path.stat().st_size > 10_000, (
        f"[{fixture_name}] PDF is only {pdf_path.stat().st_size} bytes - likely truncated"
    )

    import fitz
    doc = fitz.open(str(pdf_path))
    try:
        page_count = len(doc)
        pdf_text = "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()
    docx_text = _docx_full_text(docx_path)

    PAGE_COUNTS[fixture_name] = page_count
    record_property("page_count", page_count)
    record_property("docx_bytes", docx_path.stat().st_size)
    print(f"\n[deliverables] {fixture_name:20} PDF {page_count:>3}p "
          f"({pdf_path.stat().st_size:>9,}b)  DOCX ({docx_path.stat().st_size:>9,}b)")

    assert page_count >= 3, f"{fixture_name} rendered only {page_count} pages - build likely truncated"

    # --- Every generated action item reaches the rendered table ---
    # A dropped Things To Do row is a correctness failure, not a display
    # one: the RM works from this list, so an item that silently isn't on
    # it is an item nobody does. Counted against the rendered document,
    # not against the sort's own output, so a row lost anywhere between
    # ctx and the page is caught.
    rendered = _things_to_do_rows(docx_path)
    generated = fixture.ctx.things_to_do
    assert len(rendered) == len(generated), (
        f"[{fixture_name}] Things To Do generated {len(generated)} items but rendered "
        f"{len(rendered)} rows - an action item was dropped between the context and the page."
    )

    # ...and the rows are the ones that were generated, in deadline order,
    # renumbered 1..n. Same count with different contents would mean a row
    # was swapped rather than dropped.
    from pipeline.docx_builder import _parse_deadline
    expected_order = sorted(generated, key=lambda i: _parse_deadline(i.deadline))
    for position, (row, item) in enumerate(zip(rendered, expected_order), start=1):
        cells = [c.text.strip() for c in row.cells]
        assert cells[0] == str(position), (
            f"[{fixture_name}] row {position} is numbered {cells[0]!r} - the '#' column must "
            f"read 1..n in rendered order"
        )
        assert cells[1] == item.action and cells[2] == item.scheme, (
            f"[{fixture_name}] row {position} renders {cells[1:3]} but deadline order expects "
            f"{[item.action, item.scheme]}"
        )
        assert cells[4] == item.deadline, (
            f"[{fixture_name}] row {position} deadline is {cells[4]!r}, expected {item.deadline!r}"
        )

    # --- Parity: the same timestamp in both, from the same source ---
    expected_stamp = "Generated 16 Aug 2026 at 15:53"
    assert expected_stamp in docx_text, (
        f"[{fixture_name}] generation timestamp missing from the .docx footer"
    )
    assert expected_stamp in pdf_text, (
        f"[{fixture_name}] generation timestamp missing from the PDF footer"
    )

    # --- Parity: no section present in one output but missing from the
    # other. The PDF is rendered from the docx, so a mismatch means a
    # feature degraded in conversion rather than a content difference. ---
    from pipeline.docx_builder import SECTION_TITLES
    for title in SECTION_TITLES:
        in_docx = title in docx_text
        in_pdf = title in pdf_text
        assert in_docx == in_pdf, (
            f"[{fixture_name}] section '{title}' appears in "
            f"{'docx' if in_docx else 'PDF'} but not the other - the two deliverables have diverged."
        )

    # Invariant 12, checked against the RENDERED document - excluding the
    # named verbatim sections only (see _generated_copy_only for why).
    generated_text = _generated_copy_only(pdf_text)
    match = FORBIDDEN_RENDERED_COPY.search(generated_text)
    context = ""
    if match:
        start = max(0, match.start() - 90)
        context = generated_text[start:match.end() + 90].replace("\n", " ")
    assert match is None, (
        f"[{fixture_name}] rendered PDF contains banned copy {match.group(0)!r} in generated text - "
        f"implies a target allocation the inferred-profile model doesn't have.\n  ...{context}..."
    )


@pytest.mark.render
@requires_soffice
def test_page_counts_recorded(request):
    """Prints the collected page-count table at the end of the render run
    so an unexpected jump in any fixture is visible in one place."""
    if not PAGE_COUNTS:
        pytest.skip("render tests did not run")
    print("\n=== PDF page counts (both deliverables produced for each) ===")
    for name in FIXTURE_NAMES:
        if name in PAGE_COUNTS:
            print(f"  {name:20} {PAGE_COUNTS[name]:>3} pages")
    assert len(PAGE_COUNTS) == len(FIXTURE_NAMES), (
        f"only {len(PAGE_COUNTS)}/{len(FIXTURE_NAMES)} fixtures produced deliverables"
    )
