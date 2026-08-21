"""
tests/test_summary_client.py

Tests the one LLM call in the pipeline WITHOUT making one. Every test
here drives a fake client, so the suite runs offline and deterministically
- which is also the point being tested: the report must build with no
network at all.

What is asserted:
  * validation catches a hallucinated figure, banned target-allocation
    language, and an invented fund name;
  * the deterministic fallback renders on timeout, on network error, and
    on two consecutive validation failures;
  * build_report() raises while the summary is unapproved, and builds once
    it is approved.
"""

from __future__ import annotations

import re
from pathlib import Path

import httpx
import pytest

import anthropic

from pipeline.summary_client import (
    BANNED_PHRASES,
    MAX_ATTEMPTS,
    MODEL,
    MAX_TOKENS,
    TEMPERATURE,
    API_TIMEOUT_SECONDS,
    ClientSummary,
    build_summary_input,
    generate_client_summary,
    render_fallback_summary,
    validate_summary,
)
from tests.fixtures import ALL_FIXTURE_BUILDERS


# --------------------------------------------------------------------------
# Fakes
# --------------------------------------------------------------------------

class _Block:
    type = "text"

    def __init__(self, text):
        self.text = text


class _Usage:
    def __init__(self, i=1234, o=321):
        self.input_tokens = i
        self.output_tokens = o


class _Response:
    def __init__(self, text):
        self.content = [_Block(text)]
        self.usage = _Usage()


class FakeMessages:
    def __init__(self, replies=None, error=None):
        self._replies = list(replies or [])
        self._error = error
        self.calls = []

    def create(self, **kwargs):
        self.calls.append(kwargs)
        if self._error is not None:
            raise self._error
        return _Response(self._replies.pop(0) if self._replies else "")


class FakeClient:
    """Stands in for anthropic.Anthropic. `replies` are returned in order,
    one per attempt; `error` raises instead."""

    def __init__(self, replies=None, error=None):
        self.messages = FakeMessages(replies, error)


def _request():
    return httpx.Request("POST", "https://api.anthropic.com/v1/messages")


# --------------------------------------------------------------------------
# Fixtures
# --------------------------------------------------------------------------

@pytest.fixture(scope="module")
def golden_ctx(tmp_path_factory):
    mm = tmp_path_factory.mktemp("mm") / "mindmap.png"
    return ALL_FIXTURE_BUILDERS["A_golden"](mindmap_path=mm).ctx


@pytest.fixture(scope="module")
def payload(golden_ctx):
    return build_summary_input(golden_ctx)


# --------------------------------------------------------------------------
# The input dict
# --------------------------------------------------------------------------

def test_input_dict_carries_computed_values_only(payload):
    """Every field the summary is allowed to narrate, and nothing else -
    no raw holdings rows and no prose."""
    for key in (
        "client_name", "report_date", "total_invested_rupees", "current_value_rupees",
        "absolute_gain_rupees", "absolute_gain_pct_display",
        "value_weighted_cagr_pct_display", "cagr_method",
        "number_of_schemes", "equity_exposure_display", "risk_profile", "risk_profile_band",
        "top_holdings_by_return", "bottom_holdings_by_return", "transactions",
        "total_computed_tax_rupees", "total_computed_tax_display", "tax_zero_reason",
        "things_to_do", "earliest_deadline",
    ):
        assert key in payload, f"input dict is missing required field {key!r}"

    assert "holdings" not in payload, "raw holdings must never be sent to the model"
    assert "equity_exposure_pct" not in payload, (
        "the raw equity figure was removed - carrying 95.66 beside the display string is what "
        "let the summary print a figure the gauge renders as 95.7%"
    )
    assert "ltcg_exemption_applied_rupees" not in payload, (
        "the applied-exemption field was removed - its name led the model to describe the "
        "exempted GAIN as though it were the exemption limit"
    )
    assert payload["ltcg_annual_exemption_limit_rupees"] == 125000
    assert "not an XIRR" in payload["cagr_method"], (
        "the CAGR figure is a value-weighted average of scheme-level CAGRs; the payload must "
        "say so rather than let the summary assert more precision than the figure carries"
    )
    assert len(payload["top_holdings_by_return"]) <= 3
    assert len(payload["bottom_holdings_by_return"]) <= 3

    # Every transaction carries an action and an amount, per spec.
    for txn in payload["transactions"]:
        assert txn["action"], "transaction is missing its action"
        assert "amount_rupees" in txn


def test_equity_exposure_matches_the_single_source_function(payload, golden_ctx):
    """The summary must quote the SAME equity figure the gauge and the
    charts use - a second derivation here would be a second answer."""
    from pipeline.docx_builder import EQUITY_EXPOSURE_DECIMALS
    from pipeline.risk_profile import compute_headline_equity_exposure_pct
    value = compute_headline_equity_exposure_pct(golden_ctx.holdings)
    assert payload["equity_exposure_display"] == f"{value:.{EQUITY_EXPOSURE_DECIMALS}f}%"


# --------------------------------------------------------------------------
# Output validation
# --------------------------------------------------------------------------

def test_validation_accepts_a_faithful_summary(payload):
    text = (
        f"Your portfolio is invested across {payload['number_of_schemes']} schemes. "
        f"The current value is Rs {payload['current_value_rupees']:,.0f}. "
        f"Your equity exposure is {payload['equity_exposure_display']} and your computed "
        f"risk profile is {payload['risk_profile']}."
    )
    result = validate_summary(text, payload)
    assert result.ok, f"faithful summary was rejected: {result.reason()}"


def test_validation_catches_a_hallucinated_figure(payload):
    """The headline case: a figure that reads plausibly but was never
    computed anywhere in this report."""
    text = (
        "Your portfolio has performed well this year, delivering a gain of "
        "Rs 7,77,777 over the period."
    )
    result = validate_summary(text, payload)
    assert not result.ok
    assert any("7,77,777" in f for f in result.failures), result.failures


def test_validation_catches_a_hallucinated_percentage(payload):
    text = "Your equity exposure stands at 63.41%, which is worth discussing."
    result = validate_summary(text, payload)
    assert not result.ok
    assert any("63.41" in f for f in result.failures), result.failures


def test_validation_catches_target_allocation_language(payload):
    text = (
        "Your portfolio has moved away from its target allocation and should be "
        "brought back into line."
    )
    result = validate_summary(text, payload)
    assert not result.ok
    assert any("target allocation" in f for f in result.failures), result.failures


@pytest.mark.parametrize("phrase", BANNED_PHRASES)
def test_validation_catches_every_banned_phrase(phrase, payload):
    """Each banned phrase is caught on its own, so none of them is
    passing only because another one in the same sentence tripped."""
    result = validate_summary(f"A sentence containing {phrase} and nothing else.", payload)
    assert not result.ok
    assert any(repr(phrase) in f for f in result.failures), result.failures


def test_validation_catches_an_invented_fund_name(payload):
    text = "We suggest moving these units into the Quantum Momentum Advantage Fund."
    result = validate_summary(text, payload)
    assert not result.ok
    assert any("Quantum Momentum Advantage Fund" in f for f in result.failures), result.failures


def test_validation_allows_a_fund_name_that_is_in_the_input(payload):
    scheme = payload["transactions"][0]["scheme"]
    result = validate_summary(f"The review proposes a change to {scheme}.", payload)
    assert result.ok, f"a supplied fund name was rejected: {result.reason()}"


def test_validation_allows_standard_rupee_and_percent_formatting(payload):
    """Formatting variants of a real figure are fine; the ban is on
    figures that were never computed."""
    value = payload["current_value_rupees"]
    for written in (f"Rs {value:,.0f}", f"₹{value:,.0f}", f"{value:,.0f}", f"Rs. {value:,.0f}"):
        result = validate_summary(f"The current value is {written}.", payload)
        assert result.ok, f"{written!r} was rejected: {result.reason()}"


def test_validation_reports_every_failure_not_just_the_first(payload):
    """The retry prompt carries the reason, so it needs all of them."""
    text = (
        "Your target allocation has shifted and the portfolio is now worth "
        "Rs 9,99,99,999, so we suggest the Imaginary Growth Fund."
    )
    result = validate_summary(text, payload)
    assert not result.ok
    assert len(result.failures) >= 3, result.failures


# --------------------------------------------------------------------------
# The deterministic fallback
# --------------------------------------------------------------------------

def test_fallback_passes_its_own_validation(payload):
    """The fallback is built from the same dict, so it must satisfy the
    same checks the model's output has to satisfy. If it doesn't, the
    safety net has a hole in exactly the place it is needed most."""
    text = render_fallback_summary(payload)
    result = validate_summary(text, payload)
    assert result.ok, f"fallback failed validation: {result.reason()}"


def test_fallback_renders_for_every_fixture(tmp_path):
    """Including G_all_loss and E_no_dates, where figures are negative or
    missing - the fallback must not crash on either."""
    for name, builder in ALL_FIXTURE_BUILDERS.items():
        ctx = builder(mindmap_path=tmp_path / f"{name}.png").ctx
        data = build_summary_input(ctx)
        text = render_fallback_summary(data)
        assert len(text) > 200, f"[{name}] fallback is suspiciously short"
        result = validate_summary(text, data)
        assert result.ok, f"[{name}] fallback failed validation: {result.reason()}"


def test_fallback_used_on_timeout(payload):
    client = FakeClient(error=anthropic.APITimeoutError(request=_request()))
    summary = generate_client_summary(payload, client=client)
    assert summary.source == "fallback"
    assert summary.approved is False
    assert summary.text == render_fallback_summary(payload)
    assert any("APITimeoutError" in entry for entry in summary.failure_log), summary.failure_log


def test_fallback_used_on_network_error(payload):
    client = FakeClient(error=anthropic.APIConnectionError(request=_request()))
    summary = generate_client_summary(payload, client=client)
    assert summary.source == "fallback"
    assert summary.text == render_fallback_summary(payload)
    assert any("APIConnectionError" in entry for entry in summary.failure_log), summary.failure_log


def test_fallback_used_after_two_validation_failures(payload):
    """First reply hallucinates a figure, second uses banned language.
    Exactly two calls are made - one retry, not a loop - and the result
    is the deterministic template."""
    client = FakeClient(replies=[
        "Your portfolio gained Rs 8,88,888 over the period.",
        "Your portfolio has moved away from its target allocation.",
    ])
    summary = generate_client_summary(payload, client=client)

    assert len(client.messages.calls) == MAX_ATTEMPTS == 2
    assert summary.source == "fallback"
    assert summary.approved is False
    assert summary.text == render_fallback_summary(payload)
    assert len(summary.failure_log) == 2, summary.failure_log


def test_retry_carries_the_failure_reason_and_can_succeed(payload):
    """A first-attempt failure retries ONCE with the reason appended, and
    a valid second reply is accepted rather than discarded."""
    good = (
        f"Your portfolio is invested across {payload['number_of_schemes']} schemes and "
        f"your computed risk profile is {payload['risk_profile']}."
    )
    client = FakeClient(replies=["Your portfolio gained Rs 8,88,888.", good])
    summary = generate_client_summary(payload, client=client)

    assert summary.source == "model"
    assert summary.attempts == 2
    assert summary.text == good
    retry_prompt = client.messages.calls[1]["messages"][0]["content"]
    assert "REJECTED" in retry_prompt
    assert "8,88,888" in retry_prompt, "the retry was not told what went wrong"


def test_request_parameters_are_the_specified_ones(payload):
    client = FakeClient(replies=["Your portfolio is reviewed."])
    generate_client_summary(payload, client=client)
    call = client.messages.calls[0]
    assert call["model"] == MODEL == "claude-sonnet-4-6"
    assert call["max_tokens"] == MAX_TOKENS == 1000
    assert call["temperature"] == TEMPERATURE == 0.3
    assert API_TIMEOUT_SECONDS == 15.0
    assert "Every number in your output must appear in the input JSON" in call["system"]


def test_generated_summary_is_never_pre_approved(payload):
    client = FakeClient(replies=["Your portfolio is reviewed."])
    assert generate_client_summary(payload, client=client).approved is False
    assert ClientSummary(text="anything").approved is False


# --------------------------------------------------------------------------
# The approval gate
# --------------------------------------------------------------------------

def test_build_report_raises_when_summary_is_unapproved(golden_ctx, tmp_path):
    from pipeline.docx_builder import build_report
    golden_ctx.client_summary = ClientSummary(text="Draft summary text.", approved=False)
    try:
        with pytest.raises(ValueError, match="has not been approved"):
            build_report(golden_ctx, tmp_path / "blocked.docx")
        assert not (tmp_path / "blocked.docx").exists(), (
            "an unapproved summary must stop the build BEFORE a file is written"
        )
    finally:
        golden_ctx.client_summary = None


def test_build_report_succeeds_once_approved(golden_ctx, tmp_path):
    from pipeline.docx_builder import build_report
    text = "First paragraph of the approved summary.\n\nSecond paragraph."
    golden_ctx.client_summary = ClientSummary(text=text, approved=True)
    try:
        out = build_report(golden_ctx, tmp_path / "approved.docx")
        assert out.exists()

        from docx import Document
        body = "\n".join(p.text for p in Document(str(out)).paragraphs)
        assert "First paragraph of the approved summary." in body
        assert "Second paragraph." in body
        assert "has not been generated" not in body, (
            "the placeholder rendered even though an approved summary was present"
        )
    finally:
        golden_ctx.client_summary = None


def test_build_report_renders_placeholder_when_missing_summary_is_allowed(golden_ctx, tmp_path):
    """allow_missing_summary=True (the default until prompt 7): no summary
    is the absence of one, not an unapproved one, so the build proceeds
    and the section says so plainly."""
    from pipeline.docx_builder import build_report
    golden_ctx.client_summary = None
    assert golden_ctx.allow_missing_summary is True, "the default must stay True until prompt 7"
    out = build_report(golden_ctx, tmp_path / "placeholder.docx")
    assert out.exists()

    from docx import Document
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "has not been generated" in body


def test_build_report_raises_on_missing_summary_when_not_allowed(golden_ctx, tmp_path):
    """allow_missing_summary=False (what prompt 7 will set): a report with
    no summary at all is a build failure, not a report that ships with a
    bracketed placeholder where the client narrative should be."""
    from pipeline.docx_builder import build_report
    golden_ctx.client_summary = None
    golden_ctx.allow_missing_summary = False
    try:
        with pytest.raises(ValueError, match="Client Summary is missing"):
            build_report(golden_ctx, tmp_path / "missing.docx")
        assert not (tmp_path / "missing.docx").exists()
    finally:
        golden_ctx.allow_missing_summary = True


def test_unapproved_summary_raises_regardless_of_the_flag(golden_ctx, tmp_path):
    """allow_missing_summary governs the MISSING case only. Unapproved
    model text never renders, whatever the flag says."""
    from pipeline.docx_builder import build_report
    golden_ctx.client_summary = ClientSummary(text="Draft text.", approved=False)
    try:
        for allow in (True, False):
            golden_ctx.allow_missing_summary = allow
            with pytest.raises(ValueError, match="has not been approved"):
                build_report(golden_ctx, tmp_path / f"draft_{allow}.docx")
            assert not (tmp_path / f"draft_{allow}.docx").exists()
    finally:
        golden_ctx.client_summary = None
        golden_ctx.allow_missing_summary = True


# --------------------------------------------------------------------------
# Live API - opt-in, skipped without a credential
# --------------------------------------------------------------------------

def _has_api_credential() -> bool:
    import os
    if os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("ANTHROPIC_AUTH_TOKEN"):
        return True
    return Path(os.path.expanduser("~/.config/anthropic/credentials")).exists()


requires_api = pytest.mark.skipif(
    not _has_api_credential(),
    reason="no Anthropic credential (set ANTHROPIC_API_KEY) - live API tests are opt-in",
)


@pytest.mark.realapi
@requires_api
@pytest.mark.parametrize("fixture_name", ["A_golden", "G_all_loss"])
def test_live_api_produces_a_validated_summary(fixture_name, tmp_path, record_property):
    """One real call per fixture. G_all_loss is here deliberately: a
    portfolio down 20% is where tone and honesty matter most, and where a
    model is most tempted to reassure. The assertion is the same as the
    offline path - whatever comes back must pass validation, or the
    deterministic fallback is used instead.

    Token counts are recorded so the cost of the one LLM call in this
    pipeline stays visible on every run.
    """
    ctx = ALL_FIXTURE_BUILDERS[fixture_name](mindmap_path=tmp_path / "mm.png").ctx
    data = build_summary_input(ctx)
    summary = generate_client_summary(data)

    print(f"\n=== {fixture_name} | source={summary.source} attempts={summary.attempts} "
          f"in={summary.input_tokens} out={summary.output_tokens} ===")
    print(summary.text)
    if summary.failure_log:
        print("failure log:", summary.failure_log)

    record_property("source", summary.source)
    record_property("input_tokens", summary.input_tokens)
    record_property("output_tokens", summary.output_tokens)

    assert summary.approved is False, "live output must still arrive as an unapproved draft"
    assert validate_summary(summary.text, data).ok, (
        "the returned text must satisfy validation whether it came from the model or the fallback"
    )

    # Checks a numeric validator structurally cannot make.
    #
    # The universal-claim check applies only where the model was shown
    # FEWER schemes than the portfolio holds. On a 5-holding portfolio
    # top-3 and bottom-3 cover all five, so "all schemes are down" is
    # supported by the input and asserting against it would be wrong -
    # the dedicated 28-holding test below is where the claim is genuinely
    # an extrapolation.
    visible = {r["scheme"] for r in data["top_holdings_by_return"]}
    visible |= {r["scheme"] for r in data["bottom_holdings_by_return"]}
    if len(visible) < data["number_of_schemes"]:
        hits = _universal_quantifier_hits(summary.text, data)
        assert not hits, (
            f"universal claim {hits} while shown {len(visible)} of "
            f"{data['number_of_schemes']} schemes"
        )
    advisor = THIRD_PERSON_ADVISOR.search(summary.text)
    assert advisor is None, (
        f"refers to {advisor.group(0)!r} in the third person - this report is written by us"
    )
    assert data["equity_exposure_display"] in summary.text, (
        f"equity exposure must be stated as {data['equity_exposure_display']}, matching the gauge"
    )
    assert data["earliest_deadline"]["deadline"] in summary.text, (
        f"the earliest deadline ({data['earliest_deadline']['deadline']}) was not stated"
    )


# --------------------------------------------------------------------------
# Universal quantifiers - the class of error a numeric validator misses
# --------------------------------------------------------------------------
# The summary is shown only the top and bottom performers, never the full
# holdings list, so ANY universal claim about schemes or holdings is an
# invention even when every figure around it is correct. On the 5-holding
# fixture "all five schemes are negative" happened to be true, because
# top-3 and bottom-3 overlapped and all five were visible. On a
# 28-holding portfolio the identical sentence is unsupported - the model
# would be generalising from six rows to twenty-eight.

UNIVERSAL_QUANTIFIER = re.compile(
    r"\b(all|every|each|none|no)\b\s+"
    r"(?:of\s+)?"
    r"(?:the\s+|your\s+|these\s+|those\s+|\w+\s+)?"
    r"(schemes?|holdings?|funds?|investments?)\b",
    re.IGNORECASE,
)


def _supplied_phrases(payload: dict) -> list:
    """Phrases the payload itself hands the model. Repeating one of these
    is quoting the input, not generalising beyond it."""
    phrases = []
    for item in payload.get("things_to_do") or []:
        if item.get("what_to_do"):
            phrases.append(item["what_to_do"])
    earliest = payload.get("earliest_deadline")
    if earliest and earliest.get("what_to_do"):
        phrases.append(earliest["what_to_do"])
    return phrases


def _universal_quantifier_hits(text: str, payload: dict | None = None) -> list:
    """Universal claims about schemes or holdings that the input does not
    support.

    Payload-supplied phrasing is masked out first. "No ELSS holdings were
    found in the uploaded data" is the 80C status string quoted back - a
    statement about what the FILE contained, which the payload states
    outright, not a generalisation across holdings the model was never
    shown. Scanning the raw text flags it and would push a correct summary
    to the fallback.
    """
    scanned = text
    for phrase in _supplied_phrases(payload or {}):
        # Match the phrase loosely: the model re-words tense and articles
        # ("no ELSS holdings found in uploaded file" -> "No ELSS holdings
        # were found in the uploaded data") while keeping the key nouns.
        head = " ".join(phrase.split()[:3])
        if head:
            scanned = re.sub(re.escape(head), " ", scanned, flags=re.IGNORECASE)
    return [m.group(0) for m in UNIVERSAL_QUANTIFIER.finditer(scanned)]


@pytest.mark.parametrize("bad", [
    "All five schemes are currently showing negative returns.",
    "Every holding in the portfolio has gained this period.",
    "None of the funds are down.",
    "Each of your schemes is performing well.",
    "Returns were positive across all your holdings.",
    "No investments are showing losses.",
])
def test_universal_quantifier_detector_catches_unsupported_claims(bad):
    """Proves the detector has teeth before it is used as an assertion."""
    assert _universal_quantifier_hits(bad), f"detector missed {bad!r}"


@pytest.mark.parametrize("good", [
    "Your portfolio holds 27 schemes with a total invested amount of Rs 90,30,000.",
    "Among your holdings, HDFC Mid Cap Fund Reg (G) delivered the strongest return.",
    "Two of the schemes shown here are down over the period.",
    "The transactions proposed in this review affect three schemes.",
])
def test_universal_quantifier_detector_allows_supported_claims(good):
    """It must not fire on ordinary, properly-scoped sentences, or it
    would push every run to the fallback."""
    assert not _universal_quantifier_hits(good), f"detector false-positived on {good!r}"


def _mixed_returns_ctx(mindmap_path):
    """A_golden's 28 holdings, with the three weakest forced negative so
    the payload shows top-3 all positive and bottom-3 all negative.

    Only absolute_return_pct is touched. Current values, purchase values
    and the asset allocation are left exactly as built, so every existing
    build-time invariant (allocation vs holdings grand total, switch-pair
    balance) still holds and this stays a returns-shape change rather than
    a second, divergent fixture.
    """
    ctx = ALL_FIXTURE_BUILDERS["A_golden"](mindmap_path=mindmap_path).ctx
    ranked = sorted(
        [h for h in ctx.holdings if h.absolute_return_pct is not None],
        key=lambda h: h.absolute_return_pct,
    )
    for holding, forced in zip(ranked[:3], (-18.40, -12.75, -6.30)):
        holding.absolute_return_pct = forced
    return ctx


def test_mixed_returns_fixture_has_the_shape_the_claim_test_needs(tmp_path):
    ctx = _mixed_returns_ctx(tmp_path / "mm.png")
    data = build_summary_input(ctx)
    assert len(ctx.holdings) == 28, f"expected 28 holdings, got {len(ctx.holdings)}"
    assert len(data["top_holdings_by_return"]) == 3
    assert len(data["bottom_holdings_by_return"]) == 3
    assert all(not r["return_pct_display"].startswith("-")
               for r in data["top_holdings_by_return"]), data["top_holdings_by_return"]
    assert all(r["return_pct_display"].startswith("-")
               for r in data["bottom_holdings_by_return"]), data["bottom_holdings_by_return"]


def test_bottom_holdings_is_empty_when_nothing_is_down(tmp_path):
    """No negative returns means there is nothing to say about
    underperformers, and the payload must offer nothing to say it with."""
    ctx = ALL_FIXTURE_BUILDERS["A_golden"](mindmap_path=tmp_path / "mm.png").ctx
    for holding in ctx.holdings:
        if holding.absolute_return_pct is not None and holding.absolute_return_pct < 0:
            holding.absolute_return_pct = 4.25
    data = build_summary_input(ctx)
    assert data["bottom_holdings_by_return"] == [], data["bottom_holdings_by_return"]


@pytest.mark.realapi
@requires_api
def test_live_summary_makes_no_universal_claim_about_holdings(tmp_path, record_property):
    """The real check: 28 holdings, only six of them visible to the model,
    and no sentence in the output may generalise to the whole portfolio."""
    ctx = _mixed_returns_ctx(tmp_path / "mm.png")
    data = build_summary_input(ctx)
    summary = generate_client_summary(data)

    print(f"\n=== mixed_returns | source={summary.source} attempts={summary.attempts} "
          f"in={summary.input_tokens} out={summary.output_tokens} ===")
    print(summary.text)
    record_property("source", summary.source)

    hits = _universal_quantifier_hits(summary.text, data)
    assert not hits, (
        f"summary made a universal claim about the portfolio {hits} while being shown only "
        f"{len(data['top_holdings_by_return']) + len(data['bottom_holdings_by_return'])} of "
        f"{len(ctx.holdings)} holdings - the figures may be right but the claim is unsupported."
    )
    assert validate_summary(summary.text, data).ok


# --------------------------------------------------------------------------
# Display parity: one figure, one printed form, across the whole document
# --------------------------------------------------------------------------

def _rendered_document_text(ctx, path) -> str:
    """All user-visible text in the built .docx - body paragraphs and
    every table cell."""
    from docx import Document
    from pipeline.docx_builder import build_report
    doc = Document(str(build_report(ctx, path)))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks += [c.text for c in row.cells]
    return "\n".join(chunks)


@pytest.mark.parametrize("fixture_name", ["A_golden", "G_all_loss", "B_exemption_blown"])
def test_display_values_match_what_the_report_renders(fixture_name, tmp_path):
    """Every "_display" figure handed to the model must be the SAME string
    the report prints elsewhere.

    This is the regression that motivated the display fields: the summary
    said equity exposure was 95.66% while the risk gauge and the asset
    allocation footnote both said 95.7% - one quantity, two figures, same
    client document. Checked against the rendered .docx rather than
    against a formatter, so re-deriving a format in either place fails
    here instead of in front of a client.
    """
    ctx = ALL_FIXTURE_BUILDERS[fixture_name](mindmap_path=tmp_path / "mm.png").ctx
    payload = build_summary_input(ctx)
    rendered = _rendered_document_text(ctx, tmp_path / f"{fixture_name}.docx")

    shared = [
        "equity_exposure_display",
        "absolute_gain_pct_display",
        "value_weighted_cagr_pct_display",
        "total_invested_display",
        "current_value_display",
        "absolute_gain_display",
        "total_computed_tax_display",
    ]
    for key in shared:
        value = payload.get(key)
        if value is None:
            continue
        assert value in rendered, (
            f"[{fixture_name}] payload {key}={value!r} does not appear verbatim in the "
            f"rendered report. The summary would print a figure in a form no other section "
            f"of the same document uses."
        )

    # Per-holding figures too - these land in the Holdings Statement.
    for row in payload["top_holdings_by_return"] + payload["bottom_holdings_by_return"]:
        assert row["return_pct_display"] in rendered, (
            f"[{fixture_name}] holding return {row['return_pct_display']!r} "
            f"({row['scheme']}) is not rendered in that form"
        )
        assert row["current_value_display"] in rendered, (
            f"[{fixture_name}] holding value {row['current_value_display']!r} "
            f"({row['scheme']}) is not rendered in that form"
        )

    for txn in payload["transactions"]:
        if txn["amount_display"]:
            assert txn["amount_display"] in rendered, (
                f"[{fixture_name}] transaction amount {txn['amount_display']!r} "
                f"({txn['scheme']}) is not rendered in that form"
            )


def test_equity_display_uses_the_reports_own_precision(tmp_path):
    """The specific bug, pinned: the payload must carry 95.7%, not 95.66%."""
    from pipeline.docx_builder import EQUITY_EXPOSURE_DECIMALS
    ctx = ALL_FIXTURE_BUILDERS["A_golden"](mindmap_path=tmp_path / "mm.png").ctx
    payload = build_summary_input(ctx)
    assert EQUITY_EXPOSURE_DECIMALS == 1
    assert payload["equity_exposure_display"] == "95.7%", payload["equity_exposure_display"]


# --------------------------------------------------------------------------
# tax_zero_reason - a supported reason, so the model never invents one
# --------------------------------------------------------------------------

@pytest.mark.parametrize("fixture_name,expected", [
    ("A_golden", "gains within annual exemption"),
    ("G_all_loss", "no taxable gains - losses only"),
])
def test_tax_zero_reason_is_computed_not_inferred(fixture_name, expected, tmp_path):
    """On the all-loss fixture the model wrote that tax was zero "as the
    gains fall within the applicable thresholds" - there were no gains at
    all. Every figure was right and the causal claim was invented, which
    is precisely what a numeric validator cannot catch. The reason is now
    supplied."""
    ctx = ALL_FIXTURE_BUILDERS[fixture_name](mindmap_path=tmp_path / "mm.png").ctx
    payload = build_summary_input(ctx)
    assert payload["total_computed_tax_rupees"] == 0
    assert payload["tax_zero_reason"] == expected


def test_tax_zero_reason_is_none_when_tax_is_payable(tmp_path):
    """A non-zero tax needs no explanation, so none is offered."""
    ctx = ALL_FIXTURE_BUILDERS["B_exemption_blown"](mindmap_path=tmp_path / "mm.png").ctx
    payload = build_summary_input(ctx)
    assert payload["total_computed_tax_rupees"] > 0
    assert payload["tax_zero_reason"] is None


# --------------------------------------------------------------------------
# earliest_deadline - the item with the least time left must be stated
# --------------------------------------------------------------------------

def test_earliest_deadline_is_the_soonest_item(tmp_path):
    from pipeline.docx_builder import _parse_deadline
    for name, builder in ALL_FIXTURE_BUILDERS.items():
        ctx = builder(mindmap_path=tmp_path / f"{name}.png").ctx
        payload = build_summary_input(ctx)
        earliest = payload["earliest_deadline"]
        if not ctx.things_to_do:
            assert earliest is None
            continue
        soonest = min(_parse_deadline(i.deadline) for i in ctx.things_to_do)
        assert _parse_deadline(earliest["deadline"]) == soonest, (
            f"[{name}] earliest_deadline is {earliest['deadline']}, but the soonest item is "
            f"due {soonest}"
        )
        assert earliest["what_to_do"], "the earliest item must carry what is actually due"


def test_fallback_states_the_earliest_deadline(tmp_path):
    ctx = ALL_FIXTURE_BUILDERS["A_golden"](mindmap_path=tmp_path / "mm.png").ctx
    payload = build_summary_input(ctx)
    text = render_fallback_summary(payload)
    assert payload["earliest_deadline"]["deadline"] in text
    assert "The earliest of these is" in text


# --------------------------------------------------------------------------
# Voice: the report is written BY the firm TO the client
# --------------------------------------------------------------------------

THIRD_PERSON_ADVISOR = re.compile(
    r"\byour\s+(financial\s+)?(advisor|adviser|relationship manager|rm)\b", re.IGNORECASE
)


@pytest.mark.parametrize("bad", [
    "Please raise this with your advisor at the next meeting.",
    "Discuss the shortfall with your financial adviser.",
    "Your relationship manager will confirm the details.",
])
def test_third_person_advisor_detector_has_teeth(bad):
    assert THIRD_PERSON_ADVISOR.search(bad), f"detector missed {bad!r}"


def test_fallback_never_refers_to_the_advisor_in_third_person(tmp_path):
    """The firm is writing this. "Your advisor" would be us talking about
    ourselves as though we were somebody else."""
    for name, builder in ALL_FIXTURE_BUILDERS.items():
        ctx = builder(mindmap_path=tmp_path / f"{name}.png").ctx
        text = render_fallback_summary(build_summary_input(ctx))
        hit = THIRD_PERSON_ADVISOR.search(text)
        assert hit is None, f"[{name}] fallback says {hit.group(0)!r}"
