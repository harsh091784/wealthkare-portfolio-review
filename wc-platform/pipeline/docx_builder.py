"""
pipeline/docx_builder.py

Assembles the full WC Securities / Wealthkare Portfolio Review Report as a
.docx, in the exact 16-section order specified by the report spec.

This module ONLY builds the docx. PDF conversion and the two-pass Table of
Contents page-number resolution live in pipeline/pdf_converter.py, which
imports SECTION_TITLES and build_report() from here and calls build_report()
twice (once with placeholder TOC numbers, once with real ones detected from
the first PDF pass).

Design rules carried over from the rest of this pipeline
----------------------------------------------------------
- Never infer missing values. Where a field is genuinely absent from the
  input (no Balance Units on a SIP Stop row, no emergency-fund data, etc.)
  the cell/section renders blank or an explicit "not available" note -
  never a guessed number.
- Every dataclass here is a thin, self-contained input contract - this
  module does not re-derive business figures (risk profile, tax, NGEN
  returns); it only renders what earlier pipeline stages already computed.

Known-bug fixes applied
------------------------
- H5: every Holdings Statement table row has `w:cantSplit` set on its row
  properties, so a row is never split across a page break.
- H7: the Folio No and Balance Units columns in the Holdings Statement are
  sized generously (>=1200 / >=1000 DXA respectively) so long folio numbers
  (e.g. "477288232357") never wrap mid-digit. The Holdings table's column
  widths are set via exact-DXA integers (not a float round-trip through
  inches) summing to exactly 9020 DXA - the full A4 printable width - so
  there's no stray margin on the right and no proportional LibreOffice
  shrink pulling any column back below its H7 minimum.
- Dynamic pagination: sections no longer force a page break unconditionally.
  A lightweight running-height estimate (_LayoutCursor) decides whether a
  page break is actually needed before each section heading - only when
  the previous section's content would otherwise land within ~2 inches of
  the page bottom. Short sections (Portfolio Overview, Emergency Fund &
  Insurance Check, etc.) can now share a page with their neighbours
  instead of each getting its own mostly-blank page. The Cover Page -> TOC
  and TOC -> Director's Message transitions remain unconditional breaks,
  since those are structural conventions, not a function of content length.
"""

from __future__ import annotations

import tempfile
import warnings as _warnings
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any, Optional, Union

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

from pipeline.chart_gen import format_inr, generate_donut_chart
from pipeline.parser import Holding
from pipeline.risk_profile import AllocationRow, RiskProfileResult
from pipeline.tax_calc import PortfolioTaxResult

# --------------------------------------------------------------------------
# Brand constants
# --------------------------------------------------------------------------

NAVY = RGBColor(0x1C, 0x2B, 0x4B)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x84, 0x49)
GRAY = RGBColor(0x6E, 0x6E, 0x6E)
LIGHT_GRAY_FILL = "F0EEE6"

HEADING_SIZE = Pt(16)     # deliberately >=13pt so pdf_converter's heading
BODY_SIZE = Pt(10.5)      # detector (font-size filter) works reliably
SMALL_SIZE = Pt(9)
TABLE_HEADER_SIZE = Pt(10)
TABLE_BODY_SIZE = Pt(9.5)

FOLIO_COL_DXA = 1360   # bug H7: >=1200 DXA (measured advance width at 8pt for a 12-digit folio + buffer)
UNITS_COL_DXA = 1050   # bug H7: >=1000 DXA (measured advance width at 8pt for the "Units" header + buffer)
HOLDINGS_TABLE_FONT_SIZE = Pt(8)  # denser than the other tables' 10/9.5pt - matches the benchmark's density
HOLDINGS_TABLE_CELL_MARGINS = dict(left=40, right=40, top=8, bottom=8)  # tighter than the shared default; top/bottom only affect row height, not the column-width math
HOLDINGS_TABLE_ROW_HEIGHT_IN = 0.19  # tighter row estimate for the cursor, matching the 8pt/tight-margin sizing
DXA_PER_INCH = 1440

# Full A4 printable width at 2cm margins each side, EXACT DXA total the
# Holdings Statement table must sum to (see _set_col_widths_dxa).
PAGE_CONTENT_WIDTH_DXA = 9020


# --------------------------------------------------------------------------
# Section titles (also imported by pipeline/pdf_converter.py for TOC /
# heading detection - keep this list as the single source of truth).
# --------------------------------------------------------------------------

# The two sections whose body is copied VERBATIM from a fixed asset .docx
# (assets/director_message.docx, assets/thank_you_message.docx) instead of
# being generated by this pipeline. Named because the invariant-12
# language ban applies only to copy this pipeline writes - the Director's
# letter legitimately says "the rupee, while it drifts lower over time",
# which is currency movement, not asset-allocation drift. Anything NOT
# named here is treated as generated copy and is scanned, so adding a new
# verbatim section without adding it to VERBATIM_ASSET_SECTIONS fails the
# invariant loudly rather than being silently excluded.
DIRECTORS_MESSAGE = "Director's Message"
THANK_YOU = "Thank You"
VERBATIM_ASSET_SECTIONS = (DIRECTORS_MESSAGE, THANK_YOU)

SECTION_TITLES = [
    DIRECTORS_MESSAGE,
    "Risk Profile",
    "Portfolio Overview",
    "Current Asset Allocation",
    "Holdings Statement",
    "Mind Map",
    "Transaction Snapshot",
    "Performance Tables",
    "Tax Analysis",
    "Emergency Fund & Insurance Check",
    "Client Summary",
    "Things To Do Next",
    THANK_YOU,
    "Disclaimer",
]

assert set(VERBATIM_ASSET_SECTIONS) <= set(SECTION_TITLES)

MIND_MAP = "Mind Map"


def active_section_titles(ctx) -> list:
    """SECTION_TITLES minus any section this particular client's data has
    nothing to put in.

    Only the Mind Map is conditional. It exists to show proposed changes
    at a glance, so with no proposed transactions it has nothing to show -
    and one real client genuinely has zero actions. Printing the heading
    over "No recommended changes" invites the reader to hunt for a diagram
    that was never going to be there. The TOC is built from this same
    list, so a skipped section does not leave a dangling contents entry.
    """
    titles = list(SECTION_TITLES)
    if not ctx.transaction_snapshot:
        titles.remove(MIND_MAP)
    return titles


# --------------------------------------------------------------------------
# Dynamic pagination
# --------------------------------------------------------------------------
# python-docx has no visibility into actual rendered layout (that's decided
# by Word/LibreOffice at print time), so this is a deliberately simple
# running estimate - accurate enough to decide "is there still room for a
# short section here?" without needing a real layout engine. Every
# _build_xxx() function below feeds its own content's estimated height back
# into the shared cursor as it goes.

PAGE_HEIGHT_IN = 29.7 / 2.54
PAGE_MARGIN_IN = 2.0 / 2.54
USABLE_HEIGHT_IN = PAGE_HEIGHT_IN - 2 * PAGE_MARGIN_IN   # ~10.1in
PAGE_BREAK_THRESHOLD_IN = 2.0                             # per spec: "~2 inches"

CHARS_PER_LINE_BODY = 95
LINE_HEIGHT_IN = 0.19
PARAGRAPH_GAP_IN = 0.05
HEADING_BLOCK_HEIGHT_IN = 0.5
SPACER_HEIGHT_IN = 0.15
TABLE_HEADER_ROW_HEIGHT_IN = 0.4
TABLE_DATA_ROW_HEIGHT_IN = 0.32


class _LayoutCursor:
    """Tracks an estimated vertical position (inches) on the current page.
    `.add()` wraps modulo the usable page height so a section that
    naturally overflows onto a second (or third) page still leaves the
    cursor at a sane "how far down THIS page are we" estimate for the
    following section's break decision."""

    def __init__(self) -> None:
        self.used_in = 0.0

    def add(self, height_in: float) -> None:
        self.used_in += max(height_in, 0.0)
        if self.used_in >= USABLE_HEIGHT_IN:
            self.used_in %= USABLE_HEIGHT_IN

    def add_atomic(self, height_in: float) -> None:
        """For content that CANNOT be split across a page break - images,
        above all. `add()` models content as freely splittable, so it
        happily reports "3.0in of a 3.0in image fit in 3.1in of remaining
        space". Word/LibreOffice instead move the whole image to the next
        page, leaving that 3.1in as a visible gap the cursor never knew
        about - which is how a chart-heavy section ended up orphaning its
        charts onto a half-empty page while the cursor believed the page
        was full. Wrapping to a fresh page first keeps the estimate
        honest about where the following section actually starts."""
        height_in = max(height_in, 0.0)
        if height_in > self.remaining_in():
            self.used_in = min(height_in, USABLE_HEIGHT_IN)
        else:
            self.add(height_in)

    def remaining_in(self) -> float:
        return USABLE_HEIGHT_IN - self.used_in

    def reset(self) -> None:
        self.used_in = 0.0


def _estimate_lines(text: str, chars_per_line: int = CHARS_PER_LINE_BODY) -> int:
    if not text:
        return 1
    return max(1, -(-len(text) // chars_per_line))  # ceil division


def _measure_picture_height_in(path: Union[str, Path], width_in: float) -> float:
    """Real aspect-ratio-based height for an image inserted at width_in,
    so picture-heavy sections (Risk Profile's charts, the donut, the mind
    map) contribute an accurate estimate to the layout cursor instead of a
    guess."""
    try:
        from PIL import Image
        with Image.open(path) as img:
            w_px, h_px = img.size
        if w_px:
            return width_in * (h_px / w_px)
    except Exception:
        pass
    return width_in * 0.6  # fallback if the image can't be measured


# --------------------------------------------------------------------------
# Data model - input contract for build_report()
# --------------------------------------------------------------------------

@dataclass
class FirmInfo:
    name: str = "WC Securities Pvt Ltd"
    address: str = "A-54/A, Lower Ground Floor, Near Axis Bank, Lajpat Nagar-II, New Delhi-110024"
    arn: str = "3511"
    website: str = "https://www.wealthcareindia.com"


@dataclass
class RMInfo:
    name: str
    email: str
    phone: str


@dataclass
class PortfolioSummary:
    total_invested: float
    current_value: float
    absolute_gain: float
    absolute_gain_pct: float
    portfolio_cagr_pct: Optional[float]
    monthly_sip: float
    num_schemes: int


@dataclass
class TransactionSnapshotRow:
    scheme: str
    action: str
    amount: Optional[float]  # CURRENT value of the units being transacted (not purchase value)
    balance_units: Optional[float] = None
    purchase_amount: Optional[float] = None
    suggested_scheme: Optional[str] = None
    # Original purchase date of the units being transacted - only needed
    # (and only used) for Switch Out / Redeem rows, to classify LTCG/STCG
    # in the Tax Analysis section. See
    # pipeline.tax_calc.build_tax_holdings_from_transactions, which derives
    # the tax computation's inputs from this SAME list rather than a
    # separately-maintained holdings list, so the two sections can't drift
    # apart.
    purchase_date: Optional[date] = None
    # Amount deducted from a Switch Out's proceeds before they land in the
    # paired Switch In - capital-gains tax, and exit load if modelled.
    # Only meaningful on Switch Out rows. _validate_report_context()
    # requires Switch Out total - deductions == Switch In total, so this
    # is what makes a non-zero gap legitimate (and visible) rather than
    # money silently vanishing between the two legs.
    switch_deduction: Optional[float] = None
    switch_deduction_note: Optional[str] = None


@dataclass
class PerformanceRow:
    scheme: str
    direction: str  # "out" or "in"
    returns: dict  # {"1Y": 14.28, "CY": "N/A", ...} - percentages or "N/A"


@dataclass
class ThingsToDoRow:
    number: int
    action: str
    scheme: str
    what_to_do: str
    deadline: str
    # Lower = more urgent. This does NOT drive the rendered order -
    # _build_things_to_do() sorts by deadline ascending only, because the
    # table is a worklist read against a calendar. Priority is rendered in
    # its own visible column so importance is still communicated without
    # reordering the chronology. The default 99 means "unranked" and
    # renders as a dash rather than as the number 99.
    priority: int = 99


@dataclass
class ReportContext:
    client_name: str                       # e.g. "Rahul Sharma"
    client_salutation: str                 # "Mr." or "Ms."
    report_date: date
    firm: FirmInfo
    rm: RMInfo
    logo_path: Path

    portfolio_summary: PortfolioSummary
    asset_allocation: dict                 # {category: current_value} for the donut
    equity_sub_allocation: list            # list[AllocationRow]
    risk_profile_result: RiskProfileResult

    holdings: list                         # list[Holding], grouped/subtotalled here

    mindmap_path: Path                     # pre-rendered PNG from pipeline/mindmap.py

    transaction_snapshot: list             # list[TransactionSnapshotRow]
    performance_rows: list                 # list[PerformanceRow]
    tax_result: PortfolioTaxResult

    emergency_fund_insurance: Optional[dict]  # None -> "not available" fallback
    things_to_do: list                     # list[ThingsToDoRow]

    director_message_path: Path
    thank_you_message_path: Path

    # The generated Client Summary draft, plus its approval flag.
    # pipeline/summary_client.py builds this; the RM review screen sets
    # .approved. An unapproved summary ALWAYS fails the build, whatever
    # allow_missing_summary says - no model-written prose reaches a client
    # unreviewed.
    client_summary: Optional[Any] = None   # Optional[summary_client.ClientSummary]

    # Governs the OTHER case: no summary at all. True (the default)
    # renders client_summary_placeholder, which is what the pipeline needs
    # until prompt 7 wires the RM review screen. Prompt 7 sets this False,
    # after which a report with no summary is a build failure rather than
    # a report that quietly ships with a bracketed placeholder where the
    # client narrative should be.
    allow_missing_summary: bool = True

    client_summary_placeholder: str = (
        "[Client summary has not been generated for this run. Generate and "
        "approve it on the review screen before sending this report to a "
        "client.]"
    )

    # Stamped into the page footer of BOTH deliverables. Set once per run
    # by the caller (build_report_pdf_two_pass does this) so the docx and
    # the PDF derived from it carry the identical timestamp - if each
    # build pass called datetime.now() independently, the two-pass TOC
    # build would produce a docx and PDF stamped seconds apart, and the
    # pair would look like two different runs of the report.
    generated_at: Optional[datetime] = None


# --------------------------------------------------------------------------
# Low-level python-docx helpers
# --------------------------------------------------------------------------

def _set_row_cant_split(row) -> None:
    """Bug H5 fix: prevents a table row from splitting across a page break."""
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def _set_row_repeat_as_header(row) -> None:
    """Marks a row as a table header that REPEATS at the top of every page
    the table continues onto (`w:tblHeader`). Without this, a table that
    spans a page break shows a headerless block of numbers on the second
    page - the reader has to flip back to work out which column is which.
    Applied to every table here that can split across pages."""
    trPr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def _set_table_cell_margins(table, left: int = 60, right: int = 60, top: int = 30, bottom: int = 30) -> None:
    """Sets tight default cell margins (DXA) for the whole table. Word's
    default cell margins (~150-180 DXA / ~0.1-0.125in per side) quietly eat
    into a column's nominal width - a 1200 DXA folio column can end up with
    under 900 DXA of actual usable text space, which is exactly what caused
    a 12-digit folio number to still wrap even after sizing the column to
    the H7 minimum. Tightening margins here is part of the H7 fix, not a
    separate cosmetic change."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tblCellMar.append(el)
    tblPr.append(tblCellMar)


def _set_col_widths(table, widths_inches: list) -> None:
    """Sets column widths robustly - both the tblGrid and every cell in
    each column, since Word/LibreOffice can ignore one or the other
    depending on the autofit setting. Also tightens cell margins (see
    _set_table_cell_margins) so the declared widths translate into the
    usable text space they're supposed to."""
    table.autofit = False  # sets tblLayout to 'fixed' so explicit widths are honoured
    _set_table_cell_margins(table)
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for gridCol, width in zip(tblGrid.findall(qn("w:gridCol")), widths_inches):
            gridCol.set(qn("w:w"), str(int(width * DXA_PER_INCH)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)


def _set_col_widths_dxa(table, widths_dxa: list, cell_margins: Optional[dict] = None) -> None:
    """Sets column widths using exact DXA (twentieths of a point) integers
    - avoids the float round-trip through inches that can leave the
    tblGrid and per-cell widths a DXA or two apart, and guarantees the
    total sums to exactly what's passed in. Used for the Holdings
    Statement table, which must span exactly PAGE_CONTENT_WIDTH_DXA.

    cell_margins: optional override dict (left/right/top/bottom, in DXA)
    for _set_table_cell_margins - the Holdings table uses tighter margins
    than the default to hit its density target.
    """
    table.autofit = False
    _set_table_cell_margins(table, **(cell_margins or {}))
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for gridCol, dxa in zip(tblGrid.findall(qn("w:gridCol")), widths_dxa):
            gridCol.set(qn("w:w"), str(dxa))
    for row in table.rows:
        for cell, dxa in zip(row.cells, widths_dxa):
            cell.width = Emu(dxa * 635)  # 1 DXA == 635 EMU exactly


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _set_run(run, *, size=BODY_SIZE, bold=False, italic=False, color=None, font_name="DejaVu Sans"):
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_heading(doc: DocumentObject, text: str, cursor: _LayoutCursor, *, force_break: bool = False) -> None:
    """Section heading - deliberately rendered as a plain paragraph with an
    explicit run size (not the built-in Heading style) so the exact font
    size is guaranteed regardless of the underlying template's styles.xml.
    pdf_converter.py's TOC-page detector relies on this being consistently
    >=13pt and matching SECTION_TITLES exactly.

    force_break=True is reserved for the two structural transitions (Cover
    -> TOC, TOC -> Director's Message) that must always start a fresh page
    regardless of estimated content height. Every other section uses the
    dynamic cursor-based decision: only break if the previous section left
    less than ~2 inches of room on the current page.
    """
    if force_break:
        doc.add_page_break()
        cursor.reset()
    elif cursor.used_in > 0 and cursor.remaining_in() < PAGE_BREAK_THRESHOLD_IN:
        doc.add_page_break()
        cursor.reset()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_run(run, size=HEADING_SIZE, bold=True, color=NAVY)
    cursor.add(HEADING_BLOCK_HEIGHT_IN)


def _add_body_paragraph(doc: DocumentObject, text: str, *, size=BODY_SIZE, bold=False,
                         italic=False, color=None, align=None,
                         cursor: Optional[_LayoutCursor] = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold, italic=italic, color=color)
    if cursor is not None:
        cursor.add(_estimate_lines(text) * LINE_HEIGHT_IN + PARAGRAPH_GAP_IN)


def _add_spacer(doc: DocumentObject, cursor: _LayoutCursor) -> None:
    doc.add_paragraph()
    cursor.add(SPACER_HEIGHT_IN)


def _add_page_border(section) -> None:
    """Thin black rectangle border around every page (benchmark fix #1).
    python-docx has no high-level API for page borders - this is the
    standard `w:pgBorders` raw-XML approach, added once to the document's
    (single) section, which applies it to every page in that section."""
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")       # thin - 6 eighths-of-a-point
        el.set(qn("w:space"), "24")   # DXA offset from the page edge
        el.set(qn("w:color"), "000000")
        pgBorders.append(el)
    sectPr.append(pgBorders)


GENERATED_AT_PREFIX = "Generated"


def _generated_at_text(ctx: ReportContext) -> str:
    """Footer timestamp string, identical in the docx and the PDF derived
    from it. Falls back to now() only if the caller didn't set
    ctx.generated_at - see that field's note on why the caller should."""
    stamp = ctx.generated_at or datetime.now()
    return f"{GENERATED_AT_PREFIX} {stamp.strftime('%d %b %Y at %H:%M')}"


def _add_footer_timestamp(footer, ctx: ReportContext) -> None:
    """Appends the generation timestamp to a footer, centred and muted.

    Applied to BOTH the first-page footer and the default (all other
    pages) footer, so the stamp appears on every page of both
    deliverables - the .docx a user opens to edit and the PDF they send
    on carry the same provenance line."""
    # Reuse the footer's default empty paragraph ONLY when the footer is
    # otherwise blank. When the footer already holds content (the cover
    # page's firm/RM table), that default paragraph sits BEFORE the table,
    # so writing into it puts the timestamp above the firm block instead
    # of beneath it - append a fresh trailing paragraph in that case.
    has_content = bool(footer.tables) or any(p.text for p in footer.paragraphs)
    if has_content:
        para = footer.add_paragraph()
    else:
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    run = para.add_run(_generated_at_text(ctx))
    _set_run(run, size=Pt(7.5), color=GRAY)


DEADLINE_FORMATS = ("%d %b %Y", "%d %B %Y", "%Y-%m-%d", "%d/%m/%Y")


def _parse_deadline(deadline: str):
    """Parses a Things To Do deadline string ("30 Aug 2026") into a date.

    RAISES on anything unparseable, including a blank one. This table is
    sorted by deadline and by nothing else, so a row whose deadline can't
    be read has no defined position in it - and the failure modes of
    guessing are all bad. Sorting it last (the previous behaviour) buries
    an action item that might be the most urgent one; sorting it first
    invents urgency; and any silent handling turns a data error into a
    quietly mis-ordered worklist that nobody notices. An action item is
    correctness-critical: fail the build and fix the data.
    """
    from datetime import datetime
    raw = (deadline or "").strip()
    for fmt in DEADLINE_FORMATS:
        try:
            return datetime.strptime(raw, fmt).date()
        except (ValueError, TypeError):
            continue
    raise ValueError(
        f"Things To Do deadline {deadline!r} could not be parsed. Accepted formats: "
        f"{', '.join(DEADLINE_FORMATS)} (e.g. '30 Aug 2026'). The Things To Do table is "
        f"sorted by deadline, so a row without a readable one has no defined position - fix "
        f"the deadline on whatever built this row rather than letting the item sort to an "
        f"arbitrary place in an RM's worklist."
    )


def _format_cr(value: float) -> str:
    """Portfolio-level values in Crore format, e.g. 31552989 -> '₹3.16 Cr'."""
    return f"₹{value / 1e7:.2f} Cr"


# Precision the headline equity-exposure figure is PRINTED at, everywhere
# it appears: the risk-band line, the allocation footnote, the scatter
# annotation, and pipeline/summary_client's payload. Single source, because
# the Client Summary once said 95.66% while the gauge beside it said 95.7%.
EQUITY_EXPOSURE_DECIMALS = 1


def _format_pct(value: Optional[float], decimals: int = 2) -> str:
    if value is None:
        return ""
    return f"{value:.{decimals}f}%"


def _format_value_or_blank(value: Optional[float]) -> str:
    """Used anywhere a missing value must render as a blank cell, never a
    guess or a placeholder like 'N/A' (per the transaction-snapshot spec)."""
    if value is None:
        return ""
    return format_inr(value)


# --------------------------------------------------------------------------
# 1. Cover Page
# --------------------------------------------------------------------------

def _build_cover_page(doc: DocumentObject, ctx: ReportContext) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Vertical spacer to push the wordmark block down from the very top.
    for _ in range(3):
        doc.add_paragraph()

    # Logo LEFT of "Wealthkare" wordmark - never stacked above/below - so
    # this is a borderless 1-row, 2-column table.
    wordmark_table = doc.add_table(rows=1, cols=2)
    wordmark_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(wordmark_table)
    _set_col_widths(wordmark_table, [1.3, 4.2])

    logo_cell = wordmark_table.cell(0, 0)
    logo_cell.vertical_alignment = 1  # center
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ctx.logo_path and Path(ctx.logo_path).exists():
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(ctx.logo_path), width=Inches(1.1))

    wordmark_cell = wordmark_table.cell(0, 1)
    wordmark_para = wordmark_cell.paragraphs[0]
    wordmark_run = wordmark_para.add_run("Wealthkare")
    _set_run(wordmark_run, size=Pt(30), bold=True, color=GOLD)
    tagline_para = wordmark_cell.add_paragraph()
    tagline_run = tagline_para.add_run("Relationships Beyond Investments")
    _set_run(tagline_run, size=Pt(11), italic=True, color=NAVY)

    for _ in range(4):
        doc.add_paragraph()

    _add_body_paragraph(
        doc, "Portfolio Review Report", size=Pt(20), bold=True, color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    _add_body_paragraph(
        doc, f"{ctx.client_salutation} {ctx.client_name}", size=Pt(26), bold=True, color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    _add_body_paragraph(
        doc, f"Report as of {ctx.report_date.strftime('%d %B %Y')}", size=Pt(12), color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Two-column footer (left: firm name + address, right: RM name/email/
    # phone) - rendered into the section's FIRST-PAGE footer so it only
    # appears on the cover page.
    footer = section.first_page_footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    _remove_table_borders(footer_table)
    _set_col_widths(footer_table, [3.6, 2.9])

    left_cell = footer_table.cell(0, 0)
    left_p = left_cell.paragraphs[0]
    left_run1 = left_p.add_run(ctx.firm.name + "\n")
    _set_run(left_run1, size=SMALL_SIZE, bold=True, color=NAVY)
    left_run2 = left_p.add_run(ctx.firm.address)
    _set_run(left_run2, size=SMALL_SIZE, color=GRAY)

    right_cell = footer_table.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run1 = right_p.add_run(f"{ctx.rm.name}\n")
    _set_run(right_run1, size=SMALL_SIZE, bold=True, color=NAVY)
    right_run2 = right_p.add_run(f"{ctx.rm.email}  |  {ctx.rm.phone}")
    _set_run(right_run2, size=SMALL_SIZE, color=GRAY)

    # Generation timestamp: on the cover's own footer, and on the default
    # footer that every subsequent page uses.
    _add_footer_timestamp(footer, ctx)
    _add_footer_timestamp(section.footer, ctx)


# --------------------------------------------------------------------------
# 2. Table of Contents
# --------------------------------------------------------------------------

def _build_toc(doc: DocumentObject, toc_page_numbers: Optional[dict], cursor: _LayoutCursor,
               titles: Optional[list] = None) -> None:
    titles = titles if titles is not None else SECTION_TITLES
    # Structural transition - always its own page, regardless of how much
    # room the cover page happened to use (the cursor doesn't track the
    # cover page at all, so it can't make this decision dynamically).
    _add_heading(doc, "Table of Contents", cursor, force_break=True)
    for title in titles:
        page_str = "…"
        if toc_page_numbers is not None:
            page_str = str(toc_page_numbers.get(title, "…"))
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(f"{title}\t{page_str}")
        _set_run(run, size=Pt(12), color=NAVY)
        cursor.add(LINE_HEIGHT_IN + PARAGRAPH_GAP_IN)


# --------------------------------------------------------------------------
# 3. Director's Message / 15. Thank You - verbatim content copy
# --------------------------------------------------------------------------

def _append_verbatim_docx_body(doc: DocumentObject, source_path: Path) -> float:
    """Copies every body element (paragraphs, runs, formatting) from a
    source .docx into the target document, verbatim, via raw XML deep-copy.
    Section properties (`sectPr`) are skipped so the source file's page
    setup doesn't clobber the target's. The source assets are known to be
    plain-text letters with no embedded images (verified up front), so a
    raw element copy is sufficient to preserve formatting exactly without
    needing to also re-map image relationship IDs.

    Returns an estimated content height in inches, for the layout cursor.
    """
    source = Document(str(source_path))
    body = doc.element.body
    # Body elements must land BEFORE the target document's own w:sectPr
    # (section properties), which is always the body's last child. A plain
    # body.append() puts copied content AFTER sectPr - invalid OOXML
    # placement that LibreOffice/Word silently drop instead of rendering.
    sectPr = body.find(qn("w:sectPr"))
    total_height_in = 0.0
    for element in source.element.body:
        if element.tag == qn("w:sectPr"):
            continue
        copied = deepcopy(element)
        if sectPr is not None:
            sectPr.addprevious(copied)
        else:
            body.append(copied)

        if element.tag == qn("w:p"):
            text = "".join(t.text or "" for t in element.iter(qn("w:t")))
            total_height_in += _estimate_lines(text) * LINE_HEIGHT_IN + PARAGRAPH_GAP_IN
        elif element.tag == qn("w:tbl"):
            total_height_in += 1.0  # rough fallback - none of the current assets contain tables

    return total_height_in


def _build_directors_message(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    # Structural transition - always its own page (see _build_toc).
    _add_heading(doc, "Director's Message", cursor, force_break=True)
    height = _append_verbatim_docx_body(doc, ctx.director_message_path)
    cursor.add(height)


def _build_thank_you(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Thank You", cursor)
    height = _append_verbatim_docx_body(doc, ctx.thank_you_message_path)
    cursor.add(height)


# --------------------------------------------------------------------------
# 4. Risk Profile
# --------------------------------------------------------------------------

def _generate_risk_scatter_chart(equity_exposure_pct: Optional[float], output_path: Path) -> Path:
    """Simple scatter placeholder: the 5 risk bands plotted by their equity
    weight, with the client's actual position highlighted.

    equity_exposure_pct MUST come from risk_profile.
    compute_headline_equity_exposure_pct() - the SAME function the asset-
    allocation pie's footnote uses - so this chart's "You: X%" marker can
    never show a different number than the pie again (that was a real bug:
    91.0% on the pie vs 95.6% here, from two different calculations).
    """
    from pipeline.risk_profile import RISK_BANDS

    fig, ax = plt.subplots(figsize=(6.2, 2.8), dpi=150)
    fig.patch.set_alpha(0.0)
    ax.set_facecolor("none")

    xs = [b[1] for b in RISK_BANDS]  # equity %
    labels = [b[0] for b in RISK_BANDS]
    ax.scatter(xs, [1] * len(xs), s=90, color="#D9D5C8", zorder=2)
    # Band names are long and sit only 10 units apart on the x-axis, so a
    # plain horizontal label would overlap its neighbours - rotating them
    # keeps each one legible without needing more horizontal space.
    for x, label in zip(xs, labels):
        ax.annotate(label, (x, 1), textcoords="offset points", xytext=(4, 10),
                    ha="left", va="bottom", rotation=28, fontsize=7.5, color="#6E6E6E")

    if equity_exposure_pct is not None:
        ax.scatter([equity_exposure_pct], [1], s=170, color="#1C2B4B", zorder=3, edgecolor="#B8860B", linewidth=2)
        ax.annotate(f"You: {equity_exposure_pct:.{EQUITY_EXPOSURE_DECIMALS}f}%",
                    (equity_exposure_pct, 1), textcoords="offset points",
                    xytext=(0, -24), ha="center", fontsize=9, color="#1C2B4B", fontweight="bold")

    # Fixed range clipped the "You" marker entirely for a high-equity
    # portfolio (e.g. ~96% equity) - the point just fell outside the axes
    # with nothing drawn and no error. Extend the range to always include
    # the client's actual position, not just the 5 reference bands.
    x_min = min([30] + ([equity_exposure_pct - 8] if equity_exposure_pct is not None else []))
    x_max = max([92] + ([equity_exposure_pct + 8] if equity_exposure_pct is not None else []))
    ax.set_xlim(x_min, x_max)
    ax.set_ylim(0.55, 1.45)
    ax.get_yaxis().set_visible(False)
    ax.set_xlabel("Equity Allocation %", fontsize=9, color="#6E6E6E")
    for spine in ("top", "right", "left"):
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#D9D5C8")
    ax.tick_params(labelsize=8, colors="#6E6E6E")

    fig.savefig(output_path, dpi=200, transparent=True, bbox_inches="tight")
    plt.close(fig)
    return output_path


def _generate_traffic_light_chart(result: RiskProfileResult, output_path: Path) -> Path:
    """Simple 5-segment traffic-light bar (Conservative -> Aggressive),
    with an arrow marker over the client's current band."""
    from pipeline.risk_profile import RISK_BANDS

    labels = [b[0] for b in RISK_BANDS][::-1]  # Conservative -> Aggressive left to right
    colors = ["#C0392B", "#D98B3F", "#B8860B", "#6E9B6E", "#1E8449"]

    fig, ax = plt.subplots(figsize=(6.2, 1.7), dpi=150)
    fig.patch.set_alpha(0.0)
    ax.set_facecolor("none")

    for i, (label, color) in enumerate(zip(labels, colors)):
        ax.barh(0, 1, left=i, color=color, edgecolor="white", linewidth=2)
        # Multi-word band names don't fit a single ~1-unit-wide segment on
        # one line at a legible size - wrap onto a second line instead of
        # letting matplotlib overflow the text into neighbouring segments.
        wrapped_label = label.replace(" ", "\n", 1) if " " in label else label
        ax.text(i + 0.5, 0, wrapped_label, ha="center", va="center", fontsize=6.8,
                color="white", fontweight="bold", linespacing=1.3)

    if result.profile is not None:
        band_index = labels.index(result.profile) if result.profile in labels else None
        if band_index is not None:
            ax.annotate("▼", (band_index + 0.5, 0.62), ha="center", fontsize=14, color="#1C2B4B")

    ax.set_xlim(0, 5)
    ax.set_ylim(-0.6, 1)
    ax.axis("off")

    fig.savefig(output_path, dpi=200, transparent=True, bbox_inches="tight")
    plt.close(fig)
    return output_path


def _build_risk_profile_section(doc: DocumentObject, ctx: ReportContext, tmp_dir: Path, cursor: _LayoutCursor) -> None:
    # Forced break: this section is ~6in of largely UNSPLITTABLE content
    # (two charts, neither of which can break across a page). Letting it
    # start part-way down a page - e.g. right after the Director's Message
    # tail - pushes both charts onto the next page as a block, leaving
    # roughly half that page empty with nothing able to flow up into the
    # gap. Starting on a fresh page keeps the charts with their heading
    # and leaves enough room below them for the next short section
    # (Portfolio Overview) to be pulled up onto the same page.
    _add_heading(doc, "Risk Profile", cursor, force_break=True)
    result = ctx.risk_profile_result

    if result.profile is None:
        _add_body_paragraph(
            doc,
            "The risk profile could not be computed - no holdings with a "
            "recognised asset-class weighting were found. See the warnings "
            "log for details.",
            italic=True, color=GRAY, cursor=cursor,
        )
        for w in result.warnings:
            _add_body_paragraph(doc, f"• {w}", size=SMALL_SIZE, color=GRAY, cursor=cursor)
        return

    from pipeline.risk_profile import compute_headline_equity_exposure_pct
    equity_exposure_pct = compute_headline_equity_exposure_pct(ctx.holdings)

    # Profile name ONLY on the heading line. The band's equity range and
    # the client's own position go on a separate line below, stated as two
    # distinct facts - never an "80/20"-style ratio beside the gauge,
    # which reads as if it were the client's own split (and implies a
    # target they could deviate from; WC infers the profile from current
    # allocation, so no such target exists).
    _add_body_paragraph(doc, f"Computed Risk Profile: {result.profile}",
                         bold=True, size=Pt(13), color=NAVY, cursor=cursor)
    if result.band_definition:
        position = (
            f" This portfolio: {equity_exposure_pct:.{EQUITY_EXPOSURE_DECIMALS}f}%."
            if equity_exposure_pct is not None else ""
        )
        _add_body_paragraph(
            doc, f"{result.profile} band: {result.band_definition}.{position}",
            size=SMALL_SIZE, color=GRAY, cursor=cursor,
        )
    _add_body_paragraph(doc, result.description, cursor=cursor)
    _add_spacer(doc, cursor)

    scatter_path = tmp_dir / "risk_scatter.png"
    traffic_path = tmp_dir / "risk_traffic.png"
    _generate_risk_scatter_chart(equity_exposure_pct, scatter_path)
    _generate_traffic_light_chart(result, traffic_path)
    doc.add_picture(str(scatter_path), width=Inches(5.5))
    cursor.add_atomic(_measure_picture_height_in(scatter_path, 5.5))
    doc.add_picture(str(traffic_path), width=Inches(5.5))
    cursor.add_atomic(_measure_picture_height_in(traffic_path, 5.5))

    # The equity sub-allocation table used to render here too (duplicated
    # under Current Asset Allocation) - removed; it now lives only there.


# --------------------------------------------------------------------------
# 5. Portfolio Overview
# --------------------------------------------------------------------------

def _build_portfolio_overview(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Portfolio Overview", cursor)
    s = ctx.portfolio_summary

    rows = [
        ("Total Invested", _format_cr(s.total_invested)),
        ("Current Value", _format_cr(s.current_value)),
        ("Absolute Gain", f"{_format_cr(s.absolute_gain)}  ({_format_pct(s.absolute_gain_pct)})"),
        ("Portfolio CAGR", _format_pct(s.portfolio_cagr_pct) if s.portfolio_cagr_pct is not None else "Not available"),
        ("Monthly SIP", format_inr(s.monthly_sip)),
        ("Number of Schemes", str(s.num_schemes)),
        ("Report Date", ctx.report_date.strftime("%d %B %Y")),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    _set_col_widths(table, [2.6, 3.4])
    for label, value in rows:
        row = table.add_row().cells
        _set_run(row[0].paragraphs[0].add_run(label), size=TABLE_BODY_SIZE, bold=True, color=NAVY)
        _set_run(row[1].paragraphs[0].add_run(value), size=TABLE_BODY_SIZE)
    cursor.add(len(rows) * TABLE_DATA_ROW_HEIGHT_IN)

    # Name the CAGR method on the page - the figure is a weighted average
    # of scheme-level CAGRs, which is NOT the same thing as an XIRR, and a
    # reader would reasonably assume XIRR unless told otherwise.
    if s.portfolio_cagr_pct is not None:
        _add_body_paragraph(
            doc,
            "Portfolio CAGR: value-weighted average of scheme-level CAGRs. Not an XIRR "
            "- does not account for the timing of individual purchases.",
            size=SMALL_SIZE, italic=True, color=GRAY, cursor=cursor,
        )


# --------------------------------------------------------------------------
# 6. Current Asset Allocation
# --------------------------------------------------------------------------

def _build_asset_allocation(doc: DocumentObject, ctx: ReportContext, tmp_dir: Path, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Current Asset Allocation", cursor)

    donut_path = tmp_dir / "donut.png"
    generate_donut_chart(ctx.asset_allocation, donut_path)
    doc.add_picture(str(donut_path), width=Inches(6.0))
    cursor.add_atomic(_measure_picture_height_in(donut_path, 6.0))

    # Footnote making explicit which "equity %" definition this report
    # uses - the SAME figure the Risk Profile gauge shows (fix: the pie's
    # per-category Equity slice and the gauge's look-through figure used to
    # be two different numbers with no explanation of why).
    from pipeline.risk_profile import HEADLINE_EQUITY_EXPOSURE_FOOTNOTE, compute_headline_equity_exposure_pct
    equity_exposure_pct = compute_headline_equity_exposure_pct(ctx.holdings)
    if equity_exposure_pct is not None:
        _add_body_paragraph(
            doc,
            f"Equity exposure: {equity_exposure_pct:.{EQUITY_EXPOSURE_DECIMALS}f}%. "
            f"{HEADLINE_EQUITY_EXPOSURE_FOOTNOTE}",
            size=SMALL_SIZE, italic=True, color=GRAY, cursor=cursor,
        )

    # Market-cap sub-allocation, computed fresh from the holdings' scheme
    # names (never from a hardcoded/guessed category) - this is now the
    # ONLY place this table renders (it used to also render, duplicated,
    # under Risk Profile).
    from pipeline.risk_profile import compute_equity_market_cap_breakdown
    market_cap_rows, market_cap_warnings = compute_equity_market_cap_breakdown(ctx.holdings)
    for w in market_cap_warnings:
        _warnings.warn(f"[Current Asset Allocation] {w}", stacklevel=2)

    if market_cap_rows:
        _add_spacer(doc, cursor)
        _add_body_paragraph(doc, "Equity Sub-Allocation by Market Cap", bold=True, size=Pt(12), color=NAVY, cursor=cursor)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr_row = table.rows[0]
        _set_row_repeat_as_header(hdr_row)
        hdr = hdr_row.cells
        for cell, text in zip(hdr, ["Category", "Value", "% of Equity"]):
            _set_run(cell.paragraphs[0].add_run(text), size=TABLE_HEADER_SIZE, bold=True, color=NAVY)
        for row_data in market_cap_rows:
            row = table.add_row().cells
            label_color = RED if row_data.label == "Unclassified" else None
            _set_run(row[0].paragraphs[0].add_run(row_data.label), size=TABLE_BODY_SIZE, color=label_color)
            _set_run(row[1].paragraphs[0].add_run(format_inr(row_data.value)), size=TABLE_BODY_SIZE)
            _set_run(row[2].paragraphs[0].add_run(f"{row_data.pct_of_bucket}%"), size=TABLE_BODY_SIZE)
        cursor.add(TABLE_HEADER_ROW_HEIGHT_IN + len(market_cap_rows) * TABLE_DATA_ROW_HEIGHT_IN)
        if market_cap_warnings:
            _add_body_paragraph(
                doc,
                f"{len(market_cap_warnings)} equity scheme(s) did not match a known market-cap "
                f"keyword and are grouped under Unclassified above rather than defaulted to Large "
                f"Cap - see build warnings for the full list.",
                size=SMALL_SIZE, italic=True, color=GRAY, cursor=cursor,
            )


# --------------------------------------------------------------------------
# 7. Holdings Statement
# --------------------------------------------------------------------------

def _tight_cell_run(cell, text: str, **run_kwargs):
    """Adds a run to a table cell with paragraph spacing zeroed out - part
    of the density fix (benchmark fix #3). python-docx/Word's default
    paragraph space-before/after adds up fast across 25-30 rows."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    _set_run(run, **run_kwargs)
    return run


def _build_holdings_statement(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Holdings Statement", cursor)

    # Shortened header text, measured (via PIL font.getlength() - the true
    # advance width, not the tighter getbbox() ink extent that undercounts
    # by 20-60 DXA) to fit genuinely on ONE line at this table's 8pt font,
    # with real buffer to spare. Keeping the longer phrasing ("Balance
    # Units", "Abs. Return %") would have needed ~10470 DXA total across
    # the row to stay single-line - 1450 DXA over the 9020 budget.
    columns = ["Scheme", "Folio No", "Units", "Purchase ₹", "Current ₹", "Gain", "Return %", "CAGR %"]

    # Exact-DXA column widths summing to PAGE_CONTENT_WIDTH_DXA (9020) - the
    # full A4 printable width, so the table spans edge-to-edge with no
    # stray margin on the right - while keeping every header single-line
    # and Folio No / Units at (or above) their H7 minimums. Scheme gets
    # the remaining room after the other 7 hit their measured minimums;
    # that's enough for most real scheme names on one line, though a rare
    # very long name (e.g. "L&T Emerging Businesses Fund") wraps at a word
    # boundary rather than mid-word - acceptable, since only mid-word/
    # mid-digit splits are the defect, and Scheme was never required to be
    # single-line (only the headers were).
    #
    # Purchase ₹ / Current ₹ / Gain were previously sized only for a
    # typical HOLDING row (e.g. "₹8,30,000"). The Grand Total row's much
    # larger, BOLD aggregate figures (e.g. "₹1,08,29,790") need more room
    # than that - measured, this needs 1265 DXA vs the 980 the Current
    # column had, which is exactly why it kept wrapping after the comma.
    # Scheme and Folio No gave up the difference.
    col_widths_dxa = [1250, FOLIO_COL_DXA, UNITS_COL_DXA, 1140, 1305, 1135, 950, 830]
    assert sum(col_widths_dxa) == PAGE_CONTENT_WIDTH_DXA, (
        f"Holdings Statement column widths must sum to {PAGE_CONTENT_WIDTH_DXA} DXA, "
        f"got {sum(col_widths_dxa)}."
    )

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    _set_col_widths_dxa(table, col_widths_dxa, cell_margins=HOLDINGS_TABLE_CELL_MARGINS)

    # Header row: navy fill, white bold text - matches the benchmark.
    hdr_row = table.rows[0]
    _set_row_repeat_as_header(hdr_row)
    _set_row_cant_split(hdr_row)  # H5 - applies to every row, including header
    for cell, text in zip(hdr_row.cells, columns):
        _tight_cell_run(cell, text, size=HOLDINGS_TABLE_FONT_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, "1C2B4B")

    grouped: dict = {}
    order: list = []
    for h in ctx.holdings:
        cat = h.category or "Uncategorised"
        if cat not in grouped:
            grouped[cat] = []
            order.append(cat)
        grouped[cat].append(h)

    grand_purchase = grand_current = grand_gain = 0.0
    total_rows_rendered = 1  # header

    for cat in order:
        holdings_in_cat = grouped[cat]

        # Category divider row: light green tint, bold black text, full
        # width - matches the benchmark's separator style (fix #4). No
        # per-category subtotal row - the benchmark doesn't have one
        # either, and skipping it is part of what makes 25-30 holdings
        # fit on a single page (fix #3).
        cat_row = table.add_row()
        _set_row_cant_split(cat_row)
        cat_cell = cat_row.cells[0]
        cat_cell.merge(cat_row.cells[-1])
        _tight_cell_run(cat_cell, cat, size=HOLDINGS_TABLE_FONT_SIZE, bold=True, color=RGBColor(0x00, 0x00, 0x00))
        _shade_cell(cat_cell, "D9EAD3")
        total_rows_rendered += 1

        for h in holdings_in_cat:
            row = table.add_row()
            _set_row_cant_split(row)  # H5
            cells = row.cells
            values = [
                h.scheme or "",
                h.folio or "",
                f"{h.balance_units:,.3f}" if h.balance_units is not None else "",
                format_inr(h.purchase_value) if h.purchase_value is not None else "",
                format_inr(h.current_value) if h.current_value is not None else "",
                format_inr(h.gain) if h.gain is not None else "",
                _format_pct(h.absolute_return_pct),
                _format_pct(h.cagr_pct),
            ]
            for cell, text in zip(cells, values):
                _tight_cell_run(cell, text, size=HOLDINGS_TABLE_FONT_SIZE)
            total_rows_rendered += 1

            if h.purchase_value is not None:
                grand_purchase += h.purchase_value
            if h.current_value is not None:
                grand_current += h.current_value
            if h.gain is not None:
                grand_gain += h.gain
            elif h.purchase_value is not None and h.current_value is not None:
                grand_gain += h.current_value - h.purchase_value

    # Grand Total row: solid navy fill, white bold text, full width -
    # matches the benchmark exactly (fix #5).
    grand_row = table.add_row()
    _set_row_cant_split(grand_row)
    grand_cells = grand_row.cells
    grand_abs_return = (grand_gain / grand_purchase * 100) if grand_purchase else None
    grand_values = [
        "Grand Total", "", "",
        format_inr(grand_purchase), format_inr(grand_current), format_inr(grand_gain),
        _format_pct(grand_abs_return), "-",
    ]
    for cell, text in zip(grand_cells, grand_values):
        _tight_cell_run(cell, text, size=HOLDINGS_TABLE_FONT_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, "1C2B4B")
    total_rows_rendered += 1

    cursor.add(total_rows_rendered * HOLDINGS_TABLE_ROW_HEIGHT_IN)


# --------------------------------------------------------------------------
# 8. Mind Map
# --------------------------------------------------------------------------

def _build_mindmap_section(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Mind Map", cursor)
    _add_body_paragraph(doc, "Recommended portfolio changes at a glance:", italic=True, color=GRAY, cursor=cursor)
    if ctx.mindmap_path and Path(ctx.mindmap_path).exists():
        doc.add_picture(str(ctx.mindmap_path), width=Inches(6.5))
        cursor.add_atomic(_measure_picture_height_in(ctx.mindmap_path, 6.5))
    else:
        _add_body_paragraph(doc, "No recommended changes for this review cycle.", color=GRAY, cursor=cursor)


# --------------------------------------------------------------------------
# 9. Transaction Snapshot
# --------------------------------------------------------------------------

def _build_transaction_snapshot(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Transaction Snapshot", cursor)

    columns = ["Scheme", "Action", "Amount", "Balance Units", "Purchase Amount", "Suggested Scheme"]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    hdr_row = table.rows[0]
    _set_row_repeat_as_header(hdr_row)
    _set_row_cant_split(hdr_row)
    for cell, text in zip(hdr_row.cells, columns):
        _set_run(cell.paragraphs[0].add_run(text), size=TABLE_HEADER_SIZE, bold=True, color=NAVY)

    for r in ctx.transaction_snapshot:
        table_row = table.add_row()
        # Same treatment the Holdings and Performance tables already have -
        # without it a scheme name wrapping onto a second line can be split
        # across the page break, leaving "Kotak" on one page and "Multicap
        # Fund" orphaned at the top of the next.
        _set_row_cant_split(table_row)
        row = table_row.cells
        action_color = RED if r.action in ("Switch Out", "Redeem", "SIP Stop") else GREEN
        values = [
            r.scheme,
            r.action,
            format_inr(r.amount) if r.amount is not None else "",
            # Blank cells for missing Balance Units / Purchase Amount - never inferred.
            f"{r.balance_units:,.3f}" if r.balance_units is not None else "",
            format_inr(r.purchase_amount) if r.purchase_amount is not None else "",
            r.suggested_scheme or "",
        ]
        for i, (cell, text) in enumerate(zip(row, values)):
            run = cell.paragraphs[0].add_run(text)
            if i == 1:
                _set_run(run, size=TABLE_BODY_SIZE, bold=True, color=action_color)
            else:
                _set_run(run, size=TABLE_BODY_SIZE)

    cursor.add(TABLE_HEADER_ROW_HEIGHT_IN + len(ctx.transaction_snapshot) * TABLE_DATA_ROW_HEIGHT_IN)

    # Explicitly show any deduction taken out of a Switch Out's proceeds
    # before they reach the paired Switch In, so the gap between the two
    # legs is always accounted for on the page rather than left as an
    # unexplained difference the reader has to spot themselves.
    for t in ctx.transaction_snapshot:
        deduction = getattr(t, "switch_deduction", None)
        if t.action == "Switch Out" and deduction:
            note = getattr(t, "switch_deduction_note", None) or "capital gains tax"
            _add_body_paragraph(
                doc,
                f"{t.scheme}: switch-in amount is {format_inr(t.amount)} less "
                f"{format_inr(deduction)} ({note}) = {format_inr(t.amount - deduction)}.",
                size=SMALL_SIZE, italic=True, color=GRAY, cursor=cursor,
            )


# --------------------------------------------------------------------------
# 10. Performance Tables
# --------------------------------------------------------------------------

TRAILING_LABELS = ["1Y", "2Y", "3Y", "5Y", "7Y", "10Y", "Since Launch"]
CALENDAR_LABELS = ["CY", "CY-1", "CY-2", "CY-3", "CY-4"]

# Exact-DXA column widths, one set per table shape (8 cols for Trailing
# Returns, 6 for Calendar Year Returns - different period-column counts
# mean different leftover budgets for Scheme). Measured (PIL
# font.getlength(), true advance width) against each period column's
# longest header WORD and a realistic worst-case return value
# ("-99.99%"), with real buffer. Both leave Scheme with 2400+ DXA -
# comfortably more than any real scheme name's longest single word needs
# (e.g. "Businesses" only needs ~1310 DXA at this table's bold 9.5pt) -
# which is what actually stops LibreOffice from ever needing to fall back
# to a mid-word character break: that fallback only triggers when even a
# single word can't fit the column, and here it always can.
TRAILING_TABLE_COL_WIDTHS_DXA = [2420, 940, 940, 940, 940, 940, 940, 960]
CALENDAR_TABLE_COL_WIDTHS_DXA = [4270, 950, 950, 950, 950, 950]

assert sum(TRAILING_TABLE_COL_WIDTHS_DXA) == PAGE_CONTENT_WIDTH_DXA
assert sum(CALENDAR_TABLE_COL_WIDTHS_DXA) == PAGE_CONTENT_WIDTH_DXA


def _build_one_performance_table(
    doc: DocumentObject, title: str, labels: list, rows: list, cursor: _LayoutCursor, col_widths_dxa: list,
) -> None:
    _add_body_paragraph(doc, title, bold=True, size=Pt(12), color=NAVY, cursor=cursor)
    columns = ["Scheme"] + labels
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    _set_col_widths_dxa(table, col_widths_dxa)
    hdr_row = table.rows[0]
    _set_row_repeat_as_header(hdr_row)
    _set_row_cant_split(hdr_row)  # a row spilling across a page break was part of fix #8
    for cell, text in zip(hdr_row.cells, columns):
        _set_run(cell.paragraphs[0].add_run(text), size=TABLE_HEADER_SIZE, bold=True, color=NAVY)

    outgoing_rows = [r for r in rows if r.direction == "out"]
    incoming_rows = [r for r in rows if r.direction == "in"]

    def render_group(group_rows):
        for r in group_rows:
            row = table.add_row()
            _set_row_cant_split(row)
            cells = row.cells
            _set_run(cells[0].paragraphs[0].add_run(r.scheme), size=TABLE_BODY_SIZE, bold=True)
            for i, label in enumerate(labels, start=1):
                value = r.returns.get(label, "N/A")
                text = f"{value:.2f}%" if isinstance(value, (int, float)) else str(value)
                run = cells[i].paragraphs[0].add_run(text)
                color = RED if isinstance(value, (int, float)) and value < 0 else None
                _set_run(run, size=TABLE_BODY_SIZE, color=color)

    render_group(outgoing_rows)

    has_separator = bool(outgoing_rows and incoming_rows)
    if has_separator:
        sep_row = table.add_row()
        _set_row_cant_split(sep_row)
        for cell in sep_row.cells:
            _shade_cell(cell, "E5E0D5")
            cell.paragraphs[0].add_run("")

    render_group(incoming_rows)

    n_rows = len(rows) + (1 if has_separator else 0)
    cursor.add(TABLE_HEADER_ROW_HEIGHT_IN + n_rows * TABLE_DATA_ROW_HEIGHT_IN)


def _build_performance_tables(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Performance Tables", cursor)
    if not ctx.performance_rows:
        _add_body_paragraph(doc, "No NGEN performance data was matched for the schemes in this portfolio.", color=GRAY, cursor=cursor)
        return
    _build_one_performance_table(doc, "Trailing Returns", TRAILING_LABELS, ctx.performance_rows, cursor, TRAILING_TABLE_COL_WIDTHS_DXA)
    _add_spacer(doc, cursor)
    _build_one_performance_table(doc, "Calendar Year Returns", CALENDAR_LABELS, ctx.performance_rows, cursor, CALENDAR_TABLE_COL_WIDTHS_DXA)


# --------------------------------------------------------------------------
# 11. Tax Analysis
# --------------------------------------------------------------------------

def _build_tax_analysis(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Tax Analysis", cursor)
    s = ctx.tax_result.summary

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_row = table.rows[0]
    _set_row_repeat_as_header(hdr_row)
    hdr = hdr_row.cells
    for cell, text in zip(hdr, ["Category", "Gross Gain", "Taxable Base", "Total Tax (incl. cess)"]):
        _set_run(cell.paragraphs[0].add_run(text), size=TABLE_HEADER_SIZE, bold=True, color=NAVY)

    tax_rows = [
        ("Equity LTCG (s.112A)", s.equity_ltcg_gross_gain, s.equity_ltcg_taxable_gain, s.equity_ltcg_total_tax),
        ("Equity STCG", s.equity_stcg_gross_gain, s.equity_stcg_gross_gain, s.equity_stcg_total_tax),
        ("Non-Equity LTCG", s.non_equity_ltcg_gross_gain, s.non_equity_ltcg_gross_gain, s.non_equity_ltcg_total_tax),
        ("Non-Equity STCG", s.non_equity_stcg_gross_gain, None, None),
    ]
    for label, gross, taxable, tax in tax_rows:
        row = table.add_row().cells
        _set_run(row[0].paragraphs[0].add_run(label), size=TABLE_BODY_SIZE, bold=True)
        _set_run(row[1].paragraphs[0].add_run(format_inr(gross)), size=TABLE_BODY_SIZE)
        _set_run(row[2].paragraphs[0].add_run(format_inr(taxable) if taxable is not None else s.non_equity_stcg_note),
                 size=TABLE_BODY_SIZE, italic=(taxable is None))
        _set_run(row[3].paragraphs[0].add_run(format_inr(tax) if tax is not None else "—"), size=TABLE_BODY_SIZE)

    total_row = table.add_row().cells
    _set_run(total_row[0].paragraphs[0].add_run("Total Computed Tax"), size=TABLE_BODY_SIZE, bold=True, color=NAVY)
    _set_run(total_row[3].paragraphs[0].add_run(format_inr(s.total_computed_tax)), size=TABLE_BODY_SIZE, bold=True, color=NAVY)
    for cell in (total_row[0], total_row[1], total_row[2], total_row[3]):
        _shade_cell(cell, "E5E0D5")

    cursor.add(TABLE_HEADER_ROW_HEIGHT_IN + (len(tax_rows) + 1) * TABLE_DATA_ROW_HEIGHT_IN)

    _add_spacer(doc, cursor)
    from pipeline.tax_calc import ASSUMPTION_FLAG_TEXT
    _add_body_paragraph(doc, ASSUMPTION_FLAG_TEXT, italic=True, size=SMALL_SIZE, color=GRAY, cursor=cursor)

    # Advance-tax reminders only make sense when tax is actually payable -
    # a per-holding Q2 flag can fire on a transaction whose gain is fully
    # covered by the LTCG exemption (or offset by a loss elsewhere), which
    # left the report telling the client to pay advance tax on a Rs 0 bill.
    if s.total_computed_tax > 0:
        advance_tax_flags = sorted({r.advance_tax_flag for r in ctx.tax_result.holdings if r.advance_tax_flag})
        for flag in advance_tax_flags:
            _add_body_paragraph(doc, flag, bold=True, size=SMALL_SIZE, color=RED, cursor=cursor)

    if ctx.tax_result.tlh_opportunities:
        _add_spacer(doc, cursor)
        _add_body_paragraph(doc, "Tax-Loss Harvesting Opportunities", bold=True, size=Pt(12), color=NAVY, cursor=cursor)
        for r in ctx.tax_result.tlh_opportunities:
            # tax_calc._find_tlh_opportunities() only ever returns holdings
            # held <=12 months, so every row here IS short-term by
            # construction - say so explicitly rather than the vaguer
            # "unrealised loss", and never let it read as an LTCG item.
            _add_body_paragraph(
                doc,
                f"• {r.scheme}: short-term capital loss of {format_inr(abs(r.gain))} "
                f"(held {r.holding_period_months} months) - available for set-off against other capital gains.",
                cursor=cursor,
            )

    # The Section 80C / 80D / 80CCD block was REMOVED, not disabled.
    #
    # 80C can only be computed from ELSS purchase TRANSACTIONS dated inside
    # the current financial year. The real dashboard export carries no
    # purchase dates at all, and three of the four real client files hold
    # no ELSS to begin with, so on this data path the FY window is not
    # computable for any client - the section could only ever have printed
    # "not available in uploaded file" on every report. A row that is
    # structurally always empty is worse than no row: it reads as a gap in
    # the client's planning rather than a gap in the upload. If dated
    # purchase data becomes available, this comes back with the check that
    # made it meaningful.


# --------------------------------------------------------------------------
# 12. Emergency Fund & Insurance Check
# --------------------------------------------------------------------------

def _build_emergency_fund_insurance(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Emergency Fund & Insurance Check", cursor)
    data = ctx.emergency_fund_insurance

    if not data:
        _add_body_paragraph(
            doc,
            "Emergency fund and insurance data not available in the uploaded file. "
            "Recommend discussing with the client. Priority 1.",
            bold=True, color=RED, cursor=cursor,
        )
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_row = table.rows[0]
    _set_row_repeat_as_header(hdr_row)
    hdr = hdr_row.cells
    for cell, text in zip(hdr, ["Check", "Current", "Recommended", "Status"]):
        _set_run(cell.paragraphs[0].add_run(text), size=TABLE_HEADER_SIZE, bold=True, color=NAVY)

    def add_check_row(label, current, recommended, status_text, ok: Optional[bool]):
        row = table.add_row().cells
        _set_run(row[0].paragraphs[0].add_run(label), size=TABLE_BODY_SIZE, bold=True)
        _set_run(row[1].paragraphs[0].add_run(current), size=TABLE_BODY_SIZE)
        _set_run(row[2].paragraphs[0].add_run(recommended), size=TABLE_BODY_SIZE)
        color = GREEN if ok is True else (RED if ok is False else GRAY)
        _set_run(row[3].paragraphs[0].add_run(status_text), size=TABLE_BODY_SIZE, bold=True, color=color)

    monthly_expenses = data.get("monthly_expenses")
    liquid_assets = data.get("liquid_assets")
    if monthly_expenses is not None and liquid_assets is not None:
        target = 6 * monthly_expenses
        ok = liquid_assets >= target
        add_check_row("Emergency Fund (6x monthly expenses)", format_inr(liquid_assets), format_inr(target),
                      "Adequate" if ok else f"Shortfall of {format_inr(target - liquid_assets)}", ok)
    else:
        add_check_row("Emergency Fund (6x monthly expenses)", "not available in uploaded file", "-", "-", None)

    annual_income = data.get("annual_income")
    term_cover = data.get("existing_term_cover")
    if annual_income is not None and term_cover is not None:
        low, high = 10 * annual_income, 15 * annual_income
        ok = term_cover >= low
        add_check_row("Term Life Insurance (10-15x income)", format_inr(term_cover),
                      f"{format_inr(low)} - {format_inr(high)}",
                      "Adequate" if ok else f"Below recommended minimum of {format_inr(low)}", ok)
    else:
        add_check_row("Term Life Insurance (10-15x income)", "not available in uploaded file", "-", "-", None)

    health_cover = data.get("existing_health_cover")
    if health_cover is not None:
        target = 1_000_000
        ok = health_cover >= target
        add_check_row("Health Insurance (₹10L+ floater)", format_inr(health_cover), format_inr(target),
                      "Adequate" if ok else f"Below recommended minimum of {format_inr(target)}", ok)
    else:
        add_check_row("Health Insurance (₹10L+ floater)", "not available in uploaded file", "-", "-", None)

    cursor.add(TABLE_HEADER_ROW_HEIGHT_IN + 3 * TABLE_DATA_ROW_HEIGHT_IN)


# --------------------------------------------------------------------------
# 13. Client Summary
# --------------------------------------------------------------------------

def _build_client_summary(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Client Summary", cursor)

    summary = ctx.client_summary
    if summary is None:
        # No summary generated for this run. Rendered in grey italics so it
        # is unmistakably a placeholder rather than prose about the client.
        _add_body_paragraph(doc, ctx.client_summary_placeholder, italic=True,
                            color=GRAY, cursor=cursor)
        return

    # Reaching here means _validate_report_context() already confirmed the
    # text is approved - an unapproved summary fails the build outright
    # rather than rendering.
    for paragraph in [p.strip() for p in summary.text.split("\n\n") if p.strip()]:
        _add_body_paragraph(doc, paragraph, cursor=cursor)


# --------------------------------------------------------------------------
# 14. Things To Do Next
# --------------------------------------------------------------------------

# Measured (PIL getlength, advance width) against the bold 10pt headers
# and the one-line body values - "Tax Planning", "31 Mar 2027" - at the
# default 60/60 DXA cell margins. Must sum to the full printable width.
THINGS_TO_DO_COL_WIDTHS_DXA = [420, 1400, 1600, 3120, 1400, 1080]
assert sum(THINGS_TO_DO_COL_WIDTHS_DXA) == PAGE_CONTENT_WIDTH_DXA

# Rendered in the Priority column for a row that never set one. Printing
# the raw sentinel (99) beside real priorities of 1-4 would read as a
# genuine ranking rather than as "unranked".
UNRANKED_PRIORITY = ThingsToDoRow.__dataclass_fields__["priority"].default
UNRANKED_PRIORITY_TEXT = "-"


def _priority_text(priority: int) -> str:
    return UNRANKED_PRIORITY_TEXT if priority == UNRANKED_PRIORITY else str(priority)


def _build_things_to_do(doc: DocumentObject, ctx: ReportContext, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Things To Do Next", cursor)

    columns = ["#", "Action", "Scheme", "What to do", "Deadline", "Priority"]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    hdr_row = table.rows[0]
    _set_row_repeat_as_header(hdr_row)
    hdr = hdr_row.cells
    for cell, text in zip(hdr, columns):
        _set_run(cell.paragraphs[0].add_run(text), size=TABLE_HEADER_SIZE, bold=True, color=NAVY)

    # Sorted by DEADLINE ASCENDING, and by that alone. This table is a
    # worklist read top-down against a calendar, so the only ordering that
    # doesn't mislead is the order the dates actually fall in - an RM
    # working down the list must never meet a 31 Mar item before a 30 Aug
    # one. Priority deliberately does NOT participate in the sort; it is
    # carried in its own visible column instead, so importance is still
    # communicated without competing with the chronology. Rows are then
    # renumbered 1..n so the "#" column always matches the rendered order.
    ordered = sorted(ctx.things_to_do, key=lambda i: _parse_deadline(i.deadline))

    # Every generated item must reach the table. A dropped action item is
    # a correctness failure, not a display one - the RM works from this
    # list, and an item that silently isn't on it is an item nobody does.
    if len(ordered) != len(ctx.things_to_do):
        raise ValueError(
            f"Things To Do lost rows while ordering: {len(ctx.things_to_do)} items were "
            f"generated but only {len(ordered)} would render. Every generated item must "
            f"appear in the table."
        )

    for display_number, item in enumerate(ordered, start=1):
        row = table.add_row().cells
        values = [str(display_number), item.action, item.scheme, item.what_to_do,
                  item.deadline, _priority_text(item.priority)]
        for cell, text in zip(row, values):
            _set_run(cell.paragraphs[0].add_run(text), size=TABLE_BODY_SIZE)

    _set_col_widths_dxa(table, THINGS_TO_DO_COL_WIDTHS_DXA)

    rendered_rows = len(table.rows) - 1   # minus the header row
    if rendered_rows != len(ctx.things_to_do):
        raise ValueError(
            f"Things To Do rendered {rendered_rows} rows but {len(ctx.things_to_do)} items "
            f"were generated. Every generated action item must reach the rendered table."
        )

    cursor.add(TABLE_HEADER_ROW_HEIGHT_IN + len(ordered) * TABLE_DATA_ROW_HEIGHT_IN)


# --------------------------------------------------------------------------
# 16. Disclaimer
# --------------------------------------------------------------------------

DISCLAIMER_TEXT = (
    "Mutual Fund investments are subject to market risks. Read all scheme-related "
    "documents carefully before investing. Past performance is not indicative of "
    "future returns. This document is prepared on the basis of information "
    "provided by investors and internal sources, solely for creating awareness "
    "and educating investors. We do not accept responsibility for any investment "
    "decision taken on the basis of information provided herein. ARN: 3511 | "
    "AMFI Registered Mutual Fund & Special Investment Funds Distributor. Compare "
    "performance of similar funds at https://www.wealthcareindia.com/fund-performance/"
)


def _build_disclaimer(doc: DocumentObject, cursor: _LayoutCursor) -> None:
    _add_heading(doc, "Disclaimer", cursor)
    _add_body_paragraph(doc, DISCLAIMER_TEXT, size=SMALL_SIZE, color=GRAY, cursor=cursor)


# --------------------------------------------------------------------------
# Top-level entry point
# --------------------------------------------------------------------------

def _validate_report_context(ctx: ReportContext) -> None:
    """Fails loudly, at build time, rather than silently rendering
    inconsistent figures across sections of the same report.

    Checks the Client Summary approval gate first: text a model wrote is a
    DRAFT until a human says otherwise, so a summary that is present but
    unapproved fails the build rather than rendering. This is the last
    line of defence behind summary_client's own output validation - that
    catches hallucinated figures and banned language, this catches
    "nobody actually read it".

    Also checks: ctx.asset_allocation must sum to EXACTLY the
    holdings grand total, to the rupee. asset_allocation is meant to be
    derived purely by aggregating ctx.holdings - this assertion is what
    catches a stray hardcoded category (e.g. a "Global Equity" sliver with
    no matching holding) padding the pie beyond what the client actually
    owns, which is exactly how the pie and the Holdings Statement Grand
    Total silently disagreed in an earlier version of this pipeline.
    """
    summary = ctx.client_summary
    if summary is not None:
        if not getattr(summary, "approved", False):
            raise ValueError(
                "Client Summary has not been approved. ctx.client_summary.approved is False, so "
                "this report will not build. The summary is model-generated DRAFT text: it must "
                "be read (and edited if needed) on the RM review screen, and approved there, "
                "before it can be sent to a client. Set ctx.client_summary.approved = True only "
                "after a human has actually reviewed the text."
            )
    elif not ctx.allow_missing_summary:
        raise ValueError(
            "Client Summary is missing. ctx.client_summary is None and "
            "ctx.allow_missing_summary is False, so this report will not build. Generate the "
            "summary with pipeline.summary_client.generate_client_summary() and have it "
            "approved on the review screen. Set allow_missing_summary = True only where "
            "rendering the placeholder is genuinely intended."
        )

    holdings_total = round(sum(h.current_value for h in ctx.holdings if h.current_value is not None))
    allocation_total = round(sum(v for v in ctx.asset_allocation.values() if v is not None))
    if holdings_total != allocation_total:
        raise ValueError(
            f"Asset allocation total (Rs {allocation_total:,}) does not match the holdings "
            f"grand total (Rs {holdings_total:,}), exact to the rupee. ctx.asset_allocation must "
            f"be derived purely from ctx.holdings - check whatever built it for a stray "
            f"hardcoded category or a stale figure left over from an earlier holdings set."
        )

    # Switch pairs must balance: money leaving a Switch Out has to land in
    # its paired Switch In, less any deductions actually accounted for
    # (capital-gains tax, exit load). A previous build showed Switch Out
    # Rs 2,98,000 against Switch In Rs 2,50,000 - Rs 48,000 simply vanished
    # from the report with nothing explaining where it went.
    switch_outs = [t for t in ctx.transaction_snapshot if t.action == "Switch Out"]
    switch_ins = [t for t in ctx.transaction_snapshot if t.action == "Switch In"]
    total_switch_out = sum(t.amount for t in switch_outs if t.amount is not None)
    total_switch_in = sum(t.amount for t in switch_ins if t.amount is not None)
    total_deductions = sum(
        getattr(t, "switch_deduction", None) or 0.0 for t in switch_outs
    )
    unexplained = round(total_switch_out - total_switch_in - total_deductions)
    if switch_outs and switch_ins and unexplained != 0:
        raise ValueError(
            f"Switch pairs do not balance: Switch Out total Rs {round(total_switch_out):,} "
            f"less declared deductions Rs {round(total_deductions):,} should equal Switch In "
            f"total Rs {round(total_switch_in):,}, but Rs {unexplained:,} is unaccounted for. "
            f"Every rupee leaving a Switch Out must either land in the paired Switch In or be "
            f"explicitly declared as a deduction (tax / exit load) on the Switch Out row."
        )


def build_report(ctx: ReportContext, output_path: Union[str, Path],
                  toc_page_numbers: Optional[dict] = None) -> Path:
    """Assembles the full 16-section report and saves it as a .docx.

    toc_page_numbers: None on pass 1 (placeholder '…' shown for every TOC
    entry). On pass 2, pass in {section_title: page_number} as detected by
    pipeline/pdf_converter.py from the pass-1 PDF, to render real numbers.
    """
    _validate_report_context(ctx)

    output_path = Path(output_path)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    _add_page_border(section)

    style = doc.styles["Normal"]
    style.font.name = "DejaVu Sans"
    style.font.size = BODY_SIZE

    cursor = _LayoutCursor()

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)

        _build_cover_page(doc, ctx)                                      # 1
        _build_toc(doc, toc_page_numbers, cursor, active_section_titles(ctx))   # 2 - forced break
        _build_directors_message(doc, ctx, cursor)                       # 3 - forced break
        _build_risk_profile_section(doc, ctx, tmp_dir, cursor)           # 4 - dynamic
        _build_portfolio_overview(doc, ctx, cursor)                      # 5 - dynamic
        _build_asset_allocation(doc, ctx, tmp_dir, cursor)               # 6 - dynamic
        _build_holdings_statement(doc, ctx, cursor)                      # 7 - dynamic
        if MIND_MAP in active_section_titles(ctx):
            _build_mindmap_section(doc, ctx, cursor)                     # 8 - dynamic
        _build_transaction_snapshot(doc, ctx, cursor)                    # 9 - dynamic
        _build_performance_tables(doc, ctx, cursor)                      # 10 - dynamic
        _build_tax_analysis(doc, ctx, cursor)                            # 11 - dynamic
        _build_emergency_fund_insurance(doc, ctx, cursor)                # 12 - dynamic
        _build_client_summary(doc, ctx, cursor)                          # 13 - dynamic
        _build_things_to_do(doc, ctx, cursor)                            # 14 - dynamic
        _build_thank_you(doc, ctx, cursor)                               # 15 - dynamic
        _build_disclaimer(doc, cursor)                                   # 16 - dynamic

        doc.save(str(output_path))

    return output_path
